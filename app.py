# app.py
# ===============================
# SISTEMA SATELITAL DE ANÁLISIS AMBIENTAL INTEGRAL
# Carbono + Biodiversidad + Análisis Forrajero
# Con mapas continuos, dashboard interactivo e informe con IA (Groq)
# ===============================

# ✅ ABSOLUTAMENTE PRIMERO: Importar streamlit
import streamlit as st
# ✅ LUEGO: Configurar la página
st.set_page_config(
    page_title="Sistema Satelital de Análisis Ambiental Integral",
    page_icon="🌎",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ===== CSS MODERNO =====
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800&display=swap');

    * { font-family: 'Inter', -apple-system, sans-serif; }

    /* Fondo general */
    .stApp {
        background: linear-gradient(135deg, #0a0f1e 0%, #0d1520 50%, #0f1928 100%);
    }
    .stApp > header { background: rgba(10,15,30,0.8) !important; backdrop-filter: blur(10px); }

    /* Sidebar */
    section[data-testid="stSidebar"] {
        background: linear-gradient(180deg, #0d1520 0%, #111827 100%);
        border-right: 1px solid rgba(255,255,255,0.05);
    }
    section[data-testid="stSidebar"] .stMarkdown h1,
    section[data-testid="stSidebar"] .stMarkdown h2,
    section[data-testid="stSidebar"] .stMarkdown h3 {
        color: #f1f5f9 !important;
        font-weight: 600;
    }

    /* Tarjetas / cards */
    div.stCard, div[data-testid="stMetric"], div.stBlock {
        background: rgba(255,255,255,0.04) !important;
        backdrop-filter: blur(12px);
        border: 1px solid rgba(255,255,255,0.06);
        border-radius: 16px;
        padding: 1.25rem;
        transition: all 0.2s ease;
    }
    div.stCard:hover, div[data-testid="stMetric"]:hover {
        border-color: rgba(59,130,246,0.3);
        transform: translateY(-1px);
        box-shadow: 0 8px 30px rgba(0,0,0,0.3);
    }

    /* Métricas nativas */
    div[data-testid="stMetric"] {
        background: linear-gradient(135deg, rgba(255,255,255,0.04), rgba(255,255,255,0.01)) !important;
    }
    div[data-testid="stMetric"] > div {
        background: transparent !important;
        border: none !important;
    }
    div[data-testid="stMetric"] label {
        color: #94a3b8 !important;
        font-size: 0.8rem !important;
        font-weight: 500 !important;
        text-transform: uppercase;
        letter-spacing: 0.05em;
    }
    div[data-testid="stMetric"] div[data-testid="stMetricValue"] {
        color: #f1f5f9 !important;
        font-size: 2rem !important;
        font-weight: 700 !important;
        letter-spacing: -0.02em;
    }
    div[data-testid="stMetric"] div[data-testid="stMetricDelta"] {
        color: #10b981 !important;
        font-size: 0.75rem !important;
    }

    /* Tabs */
    .stTabs [data-baseweb="tab-list"] {
        gap: 0;
        background: rgba(255,255,255,0.03);
        border-radius: 12px;
        padding: 4px;
        border: 1px solid rgba(255,255,255,0.05);
    }
    .stTabs [data-baseweb="tab"] {
        border-radius: 8px;
        padding: 8px 18px;
        font-size: 0.85rem;
        font-weight: 500;
        color: #64748b;
        transition: all 0.15s ease;
    }
    .stTabs [data-baseweb="tab"]:hover { color: #e2e8f0; background: rgba(255,255,255,0.05); }
    .stTabs [data-baseweb="tab"][aria-selected="true"] {
        color: #f1f5f9 !important;
        background: linear-gradient(135deg, #1e3a5f 0%, #1a3a4a 100%);
        box-shadow: 0 2px 8px rgba(59,130,246,0.2);
    }

    /* Headers */
    h1, h2, h3 { color: #f1f5f9 !important; font-weight: 700; letter-spacing: -0.02em; }
    h1 { font-size: 2rem !important; background: linear-gradient(135deg, #60a5fa, #a78bfa);
         -webkit-background-clip: text; -webkit-text-fill-color: transparent; }
    h2 { font-size: 1.35rem !important; color: #e2e8f0 !important; }
    h3 { font-size: 1.1rem !important; color: #cbd5e1 !important; }

    /* Botón primario */
    .stButton button[kind="primary"] {
        background: linear-gradient(135deg, #2563eb, #7c3aed) !important;
        border: none !important;
        border-radius: 10px !important;
        font-weight: 600 !important;
        font-size: 0.9rem !important;
        padding: 0.6rem 1.2rem !important;
        transition: all 0.2s ease !important;
        box-shadow: 0 4px 15px rgba(37,99,235,0.3) !important;
    }
    .stButton button[kind="primary"]:hover {
        transform: translateY(-2px);
        box-shadow: 0 6px 25px rgba(37,99,235,0.5) !important;
    }

    /* Selectores / inputs */
    div[data-baseweb="select"], div[data-baseweb="input"] {
        background: rgba(255,255,255,0.05) !important;
        border: 1px solid rgba(255,255,255,0.1) !important;
        border-radius: 10px !important;
    }
    div[data-baseweb="select"]:hover, div[data-baseweb="input"]:hover {
        border-color: rgba(59,130,246,0.4) !important;
    }

    /* DataFrames */
    .stDataFrame { border-radius: 12px; overflow: hidden; }
    .stDataFrame table {
        background: rgba(255,255,255,0.02) !important;
        color: #e2e8f0 !important;
    }
    .stDataFrame th {
        background: rgba(59,130,246,0.15) !important;
        color: #93c5fd !important;
        font-weight: 600 !important;
        font-size: 0.75rem !important;
        text-transform: uppercase;
        letter-spacing: 0.05em;
    }
    .stDataFrame td { color: #cbd5e1 !important; border-color: rgba(255,255,255,0.05) !important; }

    /* Expander */
    .streamlit-expanderHeader {
        background: rgba(255,255,255,0.03) !important;
        border-radius: 10px !important;
        color: #e2e8f0 !important;
        font-weight: 500;
    }

    /* Info / Success / Warning boxes */
    div[data-testid="stAlert"] { border-radius: 12px !important; border: none !important; }

    /* Scrollbar */
    ::-webkit-scrollbar { width: 6px; height: 6px; }
    ::-webkit-scrollbar-track { background: rgba(255,255,255,0.02); }
    ::-webkit-scrollbar-thumb { background: rgba(255,255,255,0.15); border-radius: 3px; }
    ::-webkit-scrollbar-thumb:hover { background: rgba(255,255,255,0.25); }

    /* Footer branding */
    .footer-brand {
        text-align: center; padding: 2rem 0 1rem; font-size: 0.75rem;
        color: #475569; border-top: 1px solid rgba(255,255,255,0.04);
        margin-top: 3rem; letter-spacing: 0.02em;
    }
</style>
""", unsafe_allow_html=True)

# ===== IMPORTS ESTÁNDAR =====
import pandas as pd
import numpy as np
import tempfile
import os
import zipfile
import math
from math import log
import matplotlib.pyplot as plt
import seaborn as sns
import plotly.express as px
import plotly.graph_objects as go
import plotly.figure_factory as ff
from plotly.subplots import make_subplots
from io import BytesIO, StringIO
from datetime import datetime, timedelta
import json
import base64
import warnings
import requests
import xml.etree.ElementTree as ET
from typing import Optional, Dict, Any, List, Tuple
import random

# ✅ IMPORTACIÓN DEL MÓDULO IA (ahora con Groq)
from modules.ia_integration import (
    preparar_resumen,
    generar_analisis_carbono,
    generar_analisis_biodiversidad,
    generar_analisis_espectral,
    generar_analisis_forrajero,
    generar_recomendaciones_integradas,
    generar_plan_transicion_prv,
    available_models,
    client as groq_client,
    GROQ_API_KEY
)

# ✅ IMPORTACIÓN DEL MÓDULO PRV (Pastoreo Racional Voisin)
from modules.prv import ModeloPRV

# ===== IMPORTACIONES GOOGLE EARTH ENGINE =====
try:
    import ee
    GEE_AVAILABLE = True
except ImportError:
    GEE_AVAILABLE = False
    st.warning("⚠️ Google Earth Engine no está instalado. Para usar datos satelitales reales, instala con: pip install earthengine-api")

warnings.filterwarnings('ignore')

# ===== LIBRERÍAS GEOESPACIALES =====
import folium
from streamlit_folium import st_folium, folium_static
from folium.plugins import Fullscreen, MousePosition, HeatMap
import geopandas as gpd
from shapely.geometry import Polygon, Point, shape, MultiPolygon
from shapely.ops import unary_union
import pyproj
from branca.colormap import LinearColormap
import matplotlib.cm as cm
from scipy.interpolate import griddata
from matplotlib.colors import LinearSegmentedColormap

# ===== INICIALIZACIÓN DE GOOGLE EARTH ENGINE =====
def inicializar_gee():
    if not GEE_AVAILABLE:
        return False
    try:
        gee_secret = os.environ.get('GEE_SERVICE_ACCOUNT')
        if gee_secret:
            try:
                credentials_info = json.loads(gee_secret.strip())
                credentials = ee.ServiceAccountCredentials(
                    credentials_info['client_email'],
                    key_data=json.dumps(credentials_info)
                )
                ee.Initialize(credentials, project='ee-mawucano25')
                st.session_state.gee_authenticated = True
                st.session_state.gee_project = 'ee-mawucano25'
                return True
            except Exception as e:
                print(f"⚠️ Error con Service Account: {str(e)}")
        try:
            ee.Initialize(project='ee-mawucano25')
            st.session_state.gee_authenticated = True
            st.session_state.gee_project = 'ee-mawucano25'
            return True
        except Exception as e:
            print(f"⚠️ Error inicialización local: {str(e)}")
        st.session_state.gee_authenticated = False
        return False
    except Exception as e:
        st.session_state.gee_authenticated = False
        print(f"❌ Error crítico GEE: {str(e)}")
        return False

# ===== LIBRERÍAS PARA REPORTES =====
try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4, letter, landscape
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle,
        PageBreak, KeepTogether, PageTemplate, Frame, NextPageTemplate,
        BaseDocTemplate, FrameBreak
    )
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch, cm
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY
    from reportlab.pdfgen import canvas
    from reportlab.lib.utils import ImageReader
    REPORTPDF_AVAILABLE = True
except ImportError:
    REPORTPDF_AVAILABLE = False
    st.warning("ReportLab no está instalado. La generación de PDFs estará limitada.")

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    REPORTDOCX_AVAILABLE = True
except ImportError:
    REPORTDOCX_AVAILABLE = False
    st.warning("python-docx no está instalado. La generación de DOCX estará limitada.")

# ===============================
# 🌦️ CONECTOR CLIMÁTICO TROPICAL
# ===============================
class ConectorClimaticoTropical:
    def __init__(self):
        pass

    def obtener_datos_climaticos(self, lat: float, lon: float) -> Dict:
        if -5 <= lat <= 5 and -75 <= lon <= -50:  # Amazonía central
            return {'precipitacion': 2500 + random.uniform(-200, 200), 'temperatura': 26 + random.uniform(-1, 1)}
        elif abs(lat) < 10 and -82 <= lon <= -75:  # Chocó
            return {'precipitacion': 4000 + random.uniform(-300, 300), 'temperatura': 27 + random.uniform(-1, 1)}
        elif -15 <= lat < -5 and -70 <= lon <= -50:  # Sur amazónico
            return {'precipitacion': 1800 + random.uniform(-200, 200), 'temperatura': 25 + random.uniform(-1, 1)}
        elif -34 <= lat <= -22 and -73 <= lon <= -53:  # Argentina templada
            return {'precipitacion': 800 + random.uniform(-100, 100), 'temperatura': 18 + random.uniform(-2, 2)}
        else:
            return {'precipitacion': 1200 + random.uniform(-200, 200), 'temperatura': 22 + random.uniform(-2, 2)}

# ===============================
# 🌳 METODOLOGÍA VERRA (ajustada para cultivos y ecosistemas argentinos)
# ===============================
class MetodologiaVerra:
    def __init__(self):
        self.factores = {
            'conversion_carbono': 0.47,
            'ratio_co2': 3.67,
            'ratio_raiz': 0.24,
            'proporcion_madera_muerta': 0.15,
            'acumulacion_hojarasca': 5.0,
            'carbono_suelo': 2.5
        }
        self.factores_vegetacion = {
            # Ecosistemas originales
            'amazonia': {'factor_biomasa': 1.2, 'factor_suelo': 1.0, 'factor_madera': 1.0},
            'choco': {'factor_biomasa': 1.3, 'factor_suelo': 1.1, 'factor_madera': 1.0},
            'seco': {'factor_biomasa': 0.8, 'factor_suelo': 0.7, 'factor_madera': 0.8},
            'vid': {'factor_biomasa': 0.15, 'factor_suelo': 0.6, 'factor_madera': 0.05},
            'cultivo': {'factor_biomasa': 0.2, 'factor_suelo': 0.7, 'factor_madera': 0.1},
            'agricola': {'factor_biomasa': 0.25, 'factor_suelo': 0.8, 'factor_madera': 0.1},
            'pampa': {'factor_biomasa': 0.4, 'factor_suelo': 0.9, 'factor_madera': 0.2},
            'andes': {'factor_biomasa': 0.6, 'factor_suelo': 0.9, 'factor_madera': 0.5},
            # Nuevos ecosistemas argentinos
            'monte': {'factor_biomasa': 0.3, 'factor_suelo': 0.5, 'factor_madera': 0.3},
            'espinal': {'factor_biomasa': 0.5, 'factor_suelo': 0.8, 'factor_madera': 0.5},
            'yungas': {'factor_biomasa': 1.1, 'factor_suelo': 1.0, 'factor_madera': 1.0},
            'chaqueño': {'factor_biomasa': 0.9, 'factor_suelo': 0.9, 'factor_madera': 0.9},
            'patagonico': {'factor_biomasa': 0.3, 'factor_suelo': 0.5, 'factor_madera': 0.2},
            'paranaense': {'factor_biomasa': 1.2, 'factor_suelo': 1.1, 'factor_madera': 1.1}
        }

    def calcular_carbono_hectarea(self, ndvi: float, tipo_bosque: str, precipitacion: float) -> Dict:
        factores_veg = self.factores_vegetacion.get(tipo_bosque, 
            {'factor_biomasa': 1.0, 'factor_suelo': 1.0, 'factor_madera': 1.0})
        
        if tipo_bosque in ['vid', 'cultivo', 'agricola']:
            factor_precip = min(1.3, max(0.7, precipitacion / 1500))
        else:
            factor_precip = min(2.0, max(0.5, precipitacion / 1500))
        
        if tipo_bosque in ['vid', 'cultivo', 'agricola']:
            if ndvi > 0.7:
                agb_ton_ha = (30 + (ndvi - 0.7) * 50) * factor_precip
            elif ndvi > 0.5:
                agb_ton_ha = (20 + (ndvi - 0.5) * 60) * factor_precip
            elif ndvi > 0.3:
                agb_ton_ha = (10 + (ndvi - 0.3) * 50) * factor_precip
            else:
                agb_ton_ha = (5 + ndvi * 30) * factor_precip
        else:
            if ndvi > 0.7:
                agb_ton_ha = (150 + (ndvi - 0.7) * 300) * factor_precip
            elif ndvi > 0.5:
                agb_ton_ha = (80 + (ndvi - 0.5) * 350) * factor_precip
            elif ndvi > 0.3:
                agb_ton_ha = (30 + (ndvi - 0.3) * 250) * factor_precip
            else:
                agb_ton_ha = (5 + ndvi * 100) * factor_precip
        
        agb_ton_ha *= factores_veg['factor_biomasa']
        if tipo_bosque == "vid":
            agb_ton_ha *= 0.9
        elif tipo_bosque == "cultivo":
            agb_ton_ha *= 0.8
        
        carbono_agb = agb_ton_ha * self.factores['conversion_carbono']
        
        if tipo_bosque in ['vid', 'cultivo', 'agricola']:
            carbono_bgb = carbono_agb * (self.factores['ratio_raiz'] * 0.7)
        else:
            carbono_bgb = carbono_agb * self.factores['ratio_raiz']
        
        carbono_dw = carbono_agb * self.factores['proporcion_madera_muerta'] * factores_veg['factor_madera']
        
        if tipo_bosque in ['vid', 'cultivo', 'agricola']:
            carbono_li = self.factores['acumulacion_hojarasca'] * 0.3 * self.factores['conversion_carbono']
        else:
            carbono_li = self.factores['acumulacion_hojarasca'] * self.factores['conversion_carbono']
        
        carbono_soc = self.factores['carbono_suelo'] * factores_veg['factor_suelo']
        
        carbono_total = carbono_agb + carbono_bgb + carbono_dw + carbono_li + carbono_soc
        co2_equivalente = carbono_total * self.factores['ratio_co2']
        
        return {
            'carbono_total_ton_ha': round(carbono_total, 2),
            'co2_equivalente_ton_ha': round(co2_equivalente, 2),
            'biomasa_aerea_ton_ha': round(agb_ton_ha, 2),
            'desglose': {
                'AGB': round(carbono_agb, 2),
                'BGB': round(carbono_bgb, 2),
                'DW': round(carbono_dw, 2),
                'LI': round(carbono_li, 2),
                'SOC': round(carbono_soc, 2)
            },
            'tipo_vegetacion': tipo_bosque
        }

# ===============================
# 🦋 ANÁLISIS DE BIODIVERSIDAD (con nuevos ecosistemas)
# ===============================
class AnalisisBiodiversidad:
    def __init__(self):
        self.parametros = {
            'amazonia': {'riqueza_base': 150, 'abundancia_base': 1000, 'factor_ndvi': 0.8, 'es_cultivo': False},
            'choco': {'riqueza_base': 120, 'abundancia_base': 800, 'factor_ndvi': 0.8, 'es_cultivo': False},
            'andes': {'riqueza_base': 100, 'abundancia_base': 600, 'factor_ndvi': 0.8, 'es_cultivo': False},
            'pampa': {'riqueza_base': 50, 'abundancia_base': 300, 'factor_ndvi': 0.8, 'es_cultivo': False},
            'seco': {'riqueza_base': 40, 'abundancia_base': 200, 'factor_ndvi': 0.8, 'es_cultivo': False},
            'cultivo': {'riqueza_base': 10, 'abundancia_base': 50, 'factor_ndvi': 0.2, 'es_cultivo': True},
            'vid': {'riqueza_base': 8, 'abundancia_base': 40, 'factor_ndvi': 0.1, 'es_cultivo': True},
            'agricola': {'riqueza_base': 15, 'abundancia_base': 60, 'factor_ndvi': 0.3, 'es_cultivo': True},
            # Nuevos ecosistemas argentinos
            'monte': {'riqueza_base': 20, 'abundancia_base': 100, 'factor_ndvi': 0.5, 'es_cultivo': False},
            'espinal': {'riqueza_base': 40, 'abundancia_base': 200, 'factor_ndvi': 0.6, 'es_cultivo': False},
            'yungas': {'riqueza_base': 120, 'abundancia_base': 800, 'factor_ndvi': 0.8, 'es_cultivo': False},
            'chaqueño': {'riqueza_base': 80, 'abundancia_base': 500, 'factor_ndvi': 0.7, 'es_cultivo': False},
            'patagonico': {'riqueza_base': 20, 'abundancia_base': 150, 'factor_ndvi': 0.4, 'es_cultivo': False},
            'paranaense': {'riqueza_base': 150, 'abundancia_base': 1000, 'factor_ndvi': 0.8, 'es_cultivo': False}
        }

    def calcular_shannon(self, ndvi: float, tipo_ecosistema: str, area_ha: float, precipitacion: float) -> Dict:
        params = self.parametros.get(tipo_ecosistema, {'riqueza_base': 60, 'abundancia_base': 400, 'factor_ndvi': 0.5, 'es_cultivo': False})
        factor_ndvi = 1.0 + (ndvi * params['factor_ndvi'])
        if params['es_cultivo']:
            factor_area = min(1.3, math.log10(area_ha + 1) * 0.2 + 1)
        else:
            factor_area = min(2.0, math.log10(area_ha + 1) * 0.5 + 1)
        if tipo_ecosistema in ['amazonia', 'choco', 'yungas', 'paranaense']:
            factor_precip = min(1.5, precipitacion / 2000)
        elif params['es_cultivo']:
            factor_precip = 1.0 + (precipitacion / 2000 * 0.3)
        else:
            factor_precip = 1.0
        
        riqueza_especies = int(params['riqueza_base'] * factor_ndvi * factor_area * factor_precip * random.uniform(0.8, 1.2))
        if params['es_cultivo']:
            abundancia_total = int(params['abundancia_base'] * factor_ndvi * factor_area * factor_precip * random.uniform(0.9, 1.1))
        else:
            abundancia_total = int(params['abundancia_base'] * factor_ndvi * factor_area * factor_precip * random.uniform(0.9, 1.1))
        
        especies = []
        abundancia_acumulada = 0
        if params['es_cultivo']:
            if riqueza_especies > 0:
                abundancia_principal = int(abundancia_total * random.uniform(0.7, 0.9))
                especies.append({'especie_id': 1, 'abundancia': abundancia_principal, 'nombre': tipo_ecosistema.capitalize()})
                abundancia_acumulada += abundancia_principal
                for i in range(2, riqueza_especies + 1):
                    abundancia = int((abundancia_total - abundancia_principal) / max(riqueza_especies - 1, 1) * random.uniform(0.5, 1.5))
                    if abundancia > 0:
                        especies.append({'especie_id': i, 'abundancia': abundancia, 'nombre': f'Especie {i}'})
                        abundancia_acumulada += abundancia
        else:
            for i in range(1, riqueza_especies + 1):
                abundancia = int((abundancia_total / max(riqueza_especies, 1)) * random.lognormvariate(0, 0.5))
                if abundancia > 0:
                    especies.append({'especie_id': i, 'abundancia': abundancia, 'nombre': f'Especie {i}'})
                    abundancia_acumulada += abundancia
        
        for especie in especies:
            especie['proporcion'] = especie['abundancia'] / abundancia_acumulada if abundancia_acumulada > 0 else 0
        
        shannon = 0
        for especie in especies:
            if especie['proporcion'] > 0:
                shannon -= especie['proporcion'] * math.log(especie['proporcion'])
        
        if params['es_cultivo']:
            if shannon > 1.5:
                categoria = "Alta (para cultivo)"
                color = "#3b82f6"
            elif shannon > 1.0:
                categoria = "Moderada (para cultivo)"
                color = "#f59e0b"
            elif shannon > 0.5:
                categoria = "Baja (típico de monocultivo)"
                color = "#ef4444"
            else:
                categoria = "Muy Baja (monocultivo puro)"
                color = "#991b1b"
        else:
            if shannon > 3.5:
                categoria = "Muy Alta"
                color = "#10b981"
            elif shannon > 2.5:
                categoria = "Alta"
                color = "#3b82f6"
            elif shannon > 1.5:
                categoria = "Moderada"
                color = "#f59e0b"
            elif shannon > 0.5:
                categoria = "Baja"
                color = "#ef4444"
            else:
                categoria = "Muy Baja"
                color = "#991b1b"
        
        return {
            'indice_shannon': round(shannon, 3),
            'categoria': categoria,
            'color': color,
            'riqueza_especies': riqueza_especies,
            'abundancia_total': abundancia_acumulada,
            'especies_muestra': especies[:10],
            'es_cultivo': params['es_cultivo']
        }

# ===============================
# 🐮 ANÁLISIS FORRAJERO (completo)
# ===============================
class AnalisisForrajero:
    def __init__(self):
        self.parametros_forrajeros = {
            'pastizal_natural': {
                'productividad_kg_ms_ha': {'bajo': 2000, 'medio': 4000, 'alto': 6000},
                'eficiencia_aprovechamiento': 0.5,
                'tasa_crecimiento_diario': {'bajo': 15, 'medio': 30, 'alto': 45},
                'densidad_forraje': 2.5
            },
            'pastura_mejorada': {
                'productividad_kg_ms_ha': {'bajo': 4000, 'medio': 8000, 'alto': 12000},
                'eficiencia_aprovechamiento': 0.6,
                'tasa_crecimiento_diario': {'bajo': 25, 'medio': 50, 'alto': 75},
                'densidad_forraje': 3.0
            },
            'silvopastoril': {
                'productividad_kg_ms_ha': {'bajo': 3000, 'medio': 6000, 'alto': 9000},
                'eficiencia_aprovechamiento': 0.55,
                'tasa_crecimiento_diario': {'bajo': 20, 'medio': 40, 'alto': 60},
                'densidad_forraje': 2.8
            },
            'agroforestal': {
                'productividad_kg_ms_ha': {'bajo': 2500, 'medio': 5000, 'alto': 7500},
                'eficiencia_aprovechamiento': 0.45,
                'tasa_crecimiento_diario': {'bajo': 18, 'medio': 36, 'alto': 54},
                'densidad_forraje': 2.6
            },
            'monte': {
                'productividad_kg_ms_ha': {'bajo': 500, 'medio': 800, 'alto': 1200},
                'eficiencia_aprovechamiento': 0.3,
                'tasa_crecimiento_diario': {'bajo': 5, 'medio': 8, 'alto': 12},
                'densidad_forraje': 1.5
            },
            'patagonico': {
                'productividad_kg_ms_ha': {'bajo': 600, 'medio': 1000, 'alto': 1500},
                'eficiencia_aprovechamiento': 0.35,
                'tasa_crecimiento_diario': {'bajo': 6, 'medio': 10, 'alto': 15},
                'densidad_forraje': 1.8
            }
        }
        self.consumo_animal = {
            'vaca_adulta': 12,
            'novillo': 10,
            'ternero': 4,
            'vaca_secas': 8,
            'vaca_lactancia': 14,
            'equivalente_vaca': 12
        }
        self.factores_ndvi = {
            'bajo': {'ndvi_min': -1.0, 'ndvi_max': 0.2, 'factor': 0.3},
            'medio': {'ndvi_min': 0.2, 'ndvi_max': 0.5, 'factor': 0.6},
            'alto': {'ndvi_min': 0.5, 'ndvi_max': 1.0, 'factor': 1.0}
        }

    def estimar_disponibilidad_forrajera(self, ndvi: float, tipo_sistema: str, area_ha: float) -> Dict:
        if ndvi < 0.2:
            categoria_productividad = 'bajo'
        elif ndvi > 0.5:
            categoria_productividad = 'alto'
        else:
            categoria_productividad = 'medio'
        
        params = self.parametros_forrajeros.get(tipo_sistema, self.parametros_forrajeros['pastizal_natural'])
        productividad_base = params['productividad_kg_ms_ha'][categoria_productividad]
        factor_ndvi = 0.5 + (ndvi * 0.5)
        productividad_ajustada = productividad_base * factor_ndvi * random.uniform(0.9, 1.1)
        disponibilidad_total_kg_ms = productividad_ajustada * area_ha
        forraje_aprovechable_kg_ms = disponibilidad_total_kg_ms * params['eficiencia_aprovechamiento']
        tasa_crecimiento = params['tasa_crecimiento_diario'][categoria_productividad] * area_ha
        
        return {
            'productividad_kg_ms_ha': round(productividad_ajustada, 2),
            'disponibilidad_total_kg_ms': round(disponibilidad_total_kg_ms, 2),
            'forraje_aprovechable_kg_ms': round(forraje_aprovechable_kg_ms, 2),
            'tasa_crecimiento_diario_kg': round(tasa_crecimiento, 2),
            'categoria_productividad': categoria_productividad,
            'densidad_forraje_kg_m3': params['densidad_forraje']
        }

    def calcular_equivalentes_vaca(self, forraje_aprovechable_kg_ms: float, dias_permanencia: int = 1) -> Dict:
        consumo_ev_diario = self.consumo_animal['equivalente_vaca']
        ev_por_dia = forraje_aprovechable_kg_ms / consumo_ev_diario
        ev_para_periodo = forraje_aprovechable_kg_ms / (consumo_ev_diario * dias_permanencia)
        consumo_total_periodo = ev_para_periodo * consumo_ev_diario * dias_permanencia
        margen_seguridad = 0.8
        return {
            'ev_por_dia': round(ev_por_dia, 2),
            'ev_para_periodo': round(ev_para_periodo, 2),
            'ev_recomendado': round(ev_para_periodo * margen_seguridad, 2),
            'consumo_ev_diario_kg': consumo_ev_diario,
            'consumo_total_periodo_kg': round(consumo_total_periodo, 2),
            'dias_permanencia': dias_permanencia,
            'margen_seguridad': '20%'
        }

    def calcular_dias_permanencia(self, forraje_aprovechable_kg_ms: float, num_ev: float) -> Dict:
        consumo_ev_diario = self.consumo_animal['equivalente_vaca']
        consumo_diario_total = num_ev * consumo_ev_diario
        dias_permanencia_basico = forraje_aprovechable_kg_ms / consumo_diario_total
        dias_permanencia_ajustado = dias_permanencia_basico * 1.2
        dias_recomendados = min(30, int(dias_permanencia_ajustado))
        return {
            'dias_basico': round(dias_permanencia_basico, 1),
            'dias_ajustado': round(dias_permanencia_ajustado, 1),
            'dias_recomendados': dias_recomendados,
            'consumo_diario_total_kg': round(consumo_diario_total, 2),
            'forraje_disponible_kg': round(forraje_aprovechable_kg_ms, 2),
            'num_ev': num_ev
        }

    def dividir_lote_en_sublotes(self, area_total_ha: float, disponibilidad_forrajera_kg_ms_ha: float, heterogeneidad: float = 0.3) -> List[Dict]:
        if area_total_ha < 10:
            num_sublotes = 2
        elif area_total_ha < 50:
            num_sublotes = 3
        elif area_total_ha < 100:
            num_sublotes = 4
        else:
            num_sublotes = min(6, int(area_total_ha / 20))
        sublotes = []
        area_por_sublote = area_total_ha / num_sublotes
        for i in range(num_sublotes):
            variacion = 1 + random.uniform(-heterogeneidad, heterogeneidad)
            disponibilidad_sublote = disponibilidad_forrajera_kg_ms_ha * variacion
            forraje_sublote_kg_ms = disponibilidad_sublote * area_por_sublote
            forraje_aprovechable = forraje_sublote_kg_ms * 0.5
            sublotes.append({
                'sublote_id': i + 1,
                'area_ha': round(area_por_sublote, 2),
                'disponibilidad_kg_ms_ha': round(disponibilidad_sublote, 2),
                'forraje_total_kg_ms': round(forraje_sublote_kg_ms, 2),
                'forraje_aprovechable_kg_ms': round(forraje_aprovechable, 2),
                'productividad_relativa': round(variacion, 2)
            })
        return sublotes

    def generar_recomendaciones_rotacion(self, sublotes: List[Dict], num_ev_total: float) -> Dict:
        forraje_total_aprovechable = sum(s['forraje_aprovechable_kg_ms'] for s in sublotes)
        consumo_diario_total = num_ev_total * self.consumo_animal['equivalente_vaca']
        dias_rotacion_total = forraje_total_aprovechable / consumo_diario_total
        plan_rotacion = []
        for sublote in sublotes:
            dias_en_sublote = int((sublote['forraje_aprovechable_kg_ms'] / consumo_diario_total) * 0.8)
            dias_descanso = dias_en_sublote * 3
            plan_rotacion.append({
                'sublote': sublote['sublote_id'],
                'area_ha': sublote['area_ha'],
                'dias_uso': max(3, dias_en_sublote),
                'dias_descanso': max(21, dias_descanso),
                'productividad': sublote['productividad_relativa'],
                'recomendacion': self._generar_recomendacion_sublote(sublote['productividad_relativa'])
            })
        dias_ciclo = sum(p['dias_uso'] + p['dias_descanso'] for p in plan_rotacion) / len(plan_rotacion)
        return {
            'forraje_total_aprovechable_kg': round(forraje_total_aprovechable, 2),
            'consumo_diario_total_kg': round(consumo_diario_total, 2),
            'dias_rotacion_total': round(dias_rotacion_total, 1),
            'num_ev': num_ev_total,
            'plan_rotacion': plan_rotacion,
            'dias_ciclo_promedio': round(dias_ciclo, 1),
            'intensidad_rotacion': self._clasificar_intensidad_rotacion(dias_ciclo)
        }

    def _generar_recomendacion_sublote(self, productividad: float) -> str:
        if productividad > 1.2:
            return "Alta productividad - Considerar manejo intensivo con pastoreo rotativo"
        elif productividad > 0.8:
            return "Productividad media - Ideal para rotación estándar"
        else:
            return "Baja productividad - Requiere recuperación, considerar descanso prolongado"

    def _clasificar_intensidad_rotacion(self, dias_ciclo: float) -> str:
        if dias_ciclo < 30:
            return "Alta intensidad - Rotación rápida"
        elif dias_ciclo < 60:
            return "Media intensidad - Rotación moderada"
        else:
            return "Baja intensidad - Rotación lenta"

# ===============================
# 🗺️ SISTEMA DE MAPAS (interpolación KNN)
# ===============================
class SistemaMapas:
    SATELLITE_TILE = 'https://server.arcgisonline.com/ArcGIS/rest/services/World_Imagery/MapServer/tile/{z}/{y}/{x}'
    SATELLITE_ATTR = 'Esri, Maxar, Earthstar Geographics'
    OSM_TILE = 'OpenStreetMap'
    OSM_ATTR = 'OpenStreetMap contributors'

    def __init__(self):
        self.capa_base = self.SATELLITE_TILE
        self.estilos = {
            'area_estudio': {
                'fillColor': '#3b82f6',
                'color': '#1d4ed8',
                'weight': 4,
                'fillOpacity': 0.15,
                'dashArray': '5, 5'
            },
            'gradientes': {
                'carbono': {
                    0.0: '#0000FF', 0.2: '#00FFFF', 0.4: '#00FF00', 0.6: '#FFFF00', 0.8: '#FFA500', 1.0: '#FF0000'
                },
                'ndvi': {
                    0.0: '#8B0000', 0.2: '#FF4500', 0.4: '#FFD700', 0.6: '#9ACD32', 0.8: '#32CD32', 1.0: '#006400'
                },
                'ndwi': {
                    0.0: '#8B4513', 0.2: '#D2691E', 0.4: '#F4A460', 0.6: '#87CEEB', 0.8: '#1E90FF', 1.0: '#00008B'
                },
                'biodiversidad': {
                    0.0: '#991B1B', 0.2: '#EF4444', 0.4: '#F59E0B', 0.6: '#3B82F6', 0.8: '#8B5CF6', 1.0: '#10B981'
                },
                'forraje': {
                    0.0: '#8B4513', 0.2: '#CD853F', 0.4: '#F4A460', 0.6: '#9ACD32', 0.8: '#32CD32', 1.0: '#006400'
                },
                'ndre': {
                    0.0: '#8B0000', 0.2: '#FF4500', 0.4: '#FFD700', 0.6: '#7CFC00', 0.8: '#32CD32', 1.0: '#006400'
                },
                'msavi': {
                    0.0: '#8B4513', 0.2: '#CD853F', 0.4: '#F4A460', 0.6: '#9ACD32', 0.8: '#32CD32', 1.0: '#006400'
                },
                'evi': {
                    0.0: '#8B0000', 0.2: '#FF6347', 0.4: '#FFD700', 0.6: '#7CFC00', 0.8: '#32CD32', 1.0: '#006400'
                }
            }
        }

    @staticmethod
    def crear_mapa_con_base(gdf, zoom_extra=0):
        """Crea un mapa Folium con zoom automático al polígono."""
        bounds = gdf.total_bounds
        if any(b != b for b in bounds):
            raise ValueError("total_bounds contiene NaN")
        centro = [(bounds[1] + bounds[3]) / 2, (bounds[0] + bounds[2]) / 2]
        m = folium.Map(location=centro, zoom_start=12, tiles='OpenStreetMap', control_scale=True)
        ancho = max(bounds[2] - bounds[0], 0.001)
        alto = max(bounds[3] - bounds[1], 0.001)
        margin = max(ancho, alto) * (0.04 * (1 + zoom_extra * 0.3))
        m.fit_bounds(
            [[bounds[1] - margin, bounds[0] - margin],
             [bounds[3] + margin, bounds[2] + margin]]
        )
        Fullscreen().add_to(m)
        MousePosition().add_to(m)
        return m

    def _generar_malla_puntos(self, gdf, densidad=1200):
        if gdf is None or gdf.empty:
            return []
        try:
            poligono = gdf.geometry.iloc[0]
            bounds = gdf.total_bounds
            minx, miny, maxx, maxy = bounds
            area_ha = calcular_superficie(gdf)
            num_puntos = min(densidad, max(400, int(area_ha * 1.5)))
            puntos = []
            lado = int(np.sqrt(num_puntos))
            dx = (maxx - minx) / lado
            dy = (maxy - miny) / lado
            for i in range(lado):
                for j in range(lado):
                    lon = minx + (i + 0.5) * dx
                    lat = miny + (j + 0.5) * dy
                    punto = Point(lon, lat)
                    if poligono.contains(punto):
                        puntos.append({'lat': lat, 'lon': lon, 'x_norm': i / lado, 'y_norm': j / lado})
            return puntos
        except Exception as e:
            print(f"Error generando malla: {str(e)}")
            return []

    def _interpolar_valores_knn(self, puntos_muestra, puntos_malla, variable='carbono', k=8):
        if not puntos_muestra or not puntos_malla:
            return puntos_malla
        try:
            from sklearn.neighbors import KNeighborsRegressor
            sklearn_disponible = True
        except ImportError:
            sklearn_disponible = False

        if sklearn_disponible:
            X_train = []
            y_train = []
            for punto in puntos_muestra:
                X_train.append([punto['lat'], punto['lon']])
                if variable == 'carbono':
                    y_train.append(punto['carbono_ton_ha'])
                elif variable == 'ndvi':
                    y_train.append(punto['ndvi'])
                elif variable == 'ndwi':
                    y_train.append(punto['ndwi'])
                elif variable == 'biodiversidad':
                    y_train.append(punto['indice_shannon'])
                elif variable == 'forraje':
                    y_train.append(punto['productividad_kg_ms_ha'])
                elif variable == 'ndre':
                    y_train.append(punto['ndre'])
                elif variable == 'msavi':
                    y_train.append(punto['msavi'])
                elif variable == 'evi':
                    y_train.append(punto['evi'])

            knn = KNeighborsRegressor(n_neighbors=min(k, len(X_train)), weights='distance')
            knn.fit(X_train, y_train)
            X_pred = [[p['lat'], p['lon']] for p in puntos_malla]
            if len(X_pred) > 0:
                predicciones = knn.predict(X_pred)
                for i, punto in enumerate(puntos_malla):
                    valor = float(predicciones[i])
                    if variable == 'carbono':
                        punto['carbono_ton_ha'] = max(0, valor)
                    elif variable == 'ndvi':
                        punto['ndvi'] = max(-1.0, min(1.0, valor))
                    elif variable == 'ndwi':
                        punto['ndwi'] = max(-1.0, min(1.0, valor))
                    elif variable == 'biodiversidad':
                        punto['indice_shannon'] = max(0, valor)
                    elif variable == 'forraje':
                        punto['productividad_kg_ms_ha'] = max(0, valor)
                    elif variable == 'ndre':
                        punto['ndre'] = max(-1.0, min(1.0, valor))
                    elif variable == 'msavi':
                        punto['msavi'] = max(0, valor)
                    elif variable == 'evi':
                        punto['evi'] = max(0, valor)
        else:
            for punto_malla in puntos_malla:
                valores = []
                distancias = []
                for punto_muestra in puntos_muestra:
                    dist = np.sqrt((punto_malla['lat'] - punto_muestra['lat'])**2 + (punto_malla['lon'] - punto_muestra['lon'])**2)
                    if variable == 'carbono':
                        valor = punto_muestra['carbono_ton_ha']
                    elif variable == 'ndvi':
                        valor = punto_muestra['ndvi']
                    elif variable == 'ndwi':
                        valor = punto_muestra['ndwi']
                    elif variable == 'biodiversidad':
                        valor = punto_muestra['indice_shannon']
                    elif variable == 'forraje':
                        valor = punto_muestra['productividad_kg_ms_ha']
                    elif variable == 'ndre':
                        valor = punto_muestra['ndre']
                    elif variable == 'msavi':
                        valor = punto_muestra['msavi']
                    elif variable == 'evi':
                        valor = punto_muestra['evi']
                    peso = 1.0 / (dist ** 2) if dist > 0 else 1.0
                    valores.append(valor)
                    distancias.append(peso)
                if distancias:
                    total_pesos = sum(distancias)
                    valor_interpolado = sum(v * w for v, w in zip(valores, distancias)) / total_pesos if total_pesos > 0 else np.mean(valores)
                else:
                    valor_interpolado = 0
                if variable == 'carbono':
                    punto_malla['carbono_ton_ha'] = max(0, valor_interpolado)
                elif variable == 'ndvi':
                    punto_malla['ndvi'] = max(-1.0, min(1.0, valor_interpolado))
                elif variable == 'ndwi':
                    punto_malla['ndwi'] = max(-1.0, min(1.0, valor_interpolado))
                elif variable == 'biodiversidad':
                    punto_malla['indice_shannon'] = max(0, valor_interpolado)
                elif variable == 'forraje':
                    punto_malla['productividad_kg_ms_ha'] = max(0, valor_interpolado)
                elif variable == 'ndre':
                    punto_malla['ndre'] = max(-1.0, min(1.0, valor_interpolado))
                elif variable == 'msavi':
                    punto_malla['msavi'] = max(0, valor_interpolado)
                elif variable == 'evi':
                    punto_malla['evi'] = max(0, valor_interpolado)
        return puntos_malla

    def crear_mapa_area(self, gdf):
        if gdf is None or gdf.empty:
            return None
        try:
            m = self.crear_mapa_con_base(gdf)
            folium.GeoJson(gdf.geometry.iloc[0], style_function=lambda x: self.estilos['area_estudio'],
                           highlight_function=lambda x: {'weight': 6, 'color': '#1e40af', 'fillOpacity': 0.3}).add_to(m)
            return m
        except Exception as e:
            st.warning(f"Error al crear mapa: {str(e)}")
            return None

    def crear_mapa_calor_interpolado(self, resultados, variable='carbono', gdf_area=None):
        if not resultados or gdf_area is None or gdf_area.empty:
            return None
        try:
            puntos_muestra = resultados.get(f'puntos_{variable}', [])
            if not puntos_muestra:
                return None
            puntos_malla = self._generar_malla_puntos(gdf_area, densidad=1200)
            if not puntos_malla:
                return None
            puntos_interpolados = self._interpolar_valores_knn(puntos_muestra, puntos_malla, variable)
            bounds = gdf_area.total_bounds
            centro = [(bounds[1] + bounds[3]) / 2, (bounds[0] + bounds[2]) / 2]
            m = self.crear_mapa_con_base(gdf_area, zoom_extra=2)
            folium.GeoJson(gdf_area.geometry.iloc[0], style_function=lambda x: {
                'fillColor': 'transparent', 'color': '#1d4ed8', 'weight': 2, 'fillOpacity': 0.05, 'dashArray': '5, 5'
            }).add_to(m)
            heat_data = []
            for punto in puntos_interpolados:
                if variable == 'carbono':
                    heat_data.append([punto['lat'], punto['lon'], punto['carbono_ton_ha']])
                elif variable == 'ndvi':
                    heat_data.append([punto['lat'], punto['lon'], punto['ndvi']])
                elif variable == 'ndwi':
                    heat_data.append([punto['lat'], punto['lon'], punto['ndwi']])
                elif variable == 'biodiversidad':
                    heat_data.append([punto['lat'], punto['lon'], punto['indice_shannon']])
                elif variable == 'forraje':
                    heat_data.append([punto['lat'], punto['lon'], punto['productividad_kg_ms_ha']])
                elif variable == 'ndre':
                    heat_data.append([punto['lat'], punto['lon'], punto['ndre']])
                elif variable == 'msavi':
                    heat_data.append([punto['lat'], punto['lon'], punto['msavi']])
                elif variable == 'evi':
                    heat_data.append([punto['lat'], punto['lon'], punto['evi']])
            gradient = self.estilos['gradientes'].get(variable, self.estilos['gradientes']['carbono'])
            radius = 45 if variable in ['carbono', 'biodiversidad', 'forraje'] else 40
            blur = 40 if variable in ['carbono', 'biodiversidad', 'forraje'] else 35
            HeatMap(heat_data, name=variable, min_opacity=0.7, radius=radius, blur=blur, gradient=gradient, max_zoom=18).add_to(m)
            return m
        except Exception as e:
            st.warning(f"Error al crear mapa de calor para {variable}: {str(e)}")
            return None

    def crear_mapa_combinado_interpolado(self, resultados, gdf_area=None):
        """
        Crea un mapa con múltiples capas de calor continuas (carbono, ndvi, ndwi, biodiversidad, forraje)
        y control de capas para activar/desactivar cada una.
        """
        if not resultados or gdf_area is None or gdf_area.empty:
            return None
        try:
            bounds = gdf_area.total_bounds
            centro = [(bounds[1] + bounds[3]) / 2, (bounds[0] + bounds[2]) / 2]
            m = self.crear_mapa_con_base(gdf_area, zoom_extra=2)
            # Capa base: polígono transparente
            folium.GeoJson(gdf_area.geometry.iloc[0], style_function=lambda x: {
                'fillColor': 'transparent', 'color': '#1d4ed8', 'weight': 2,
                'fillOpacity': 0.05, 'dashArray': '5, 5'
            }).add_to(m)

            # Variables a incluir y su configuración
            variables = [
                ('carbono', '🌳 Carbono', 45, 40, False),
                ('ndvi', '📈 NDVI', 40, 35, False),
                ('ndwi', '💧 NDWI', 40, 35, False),
                ('biodiversidad', '🦋 Biodiversidad', 45, 40, False),
                ('forraje', '🌿 Forraje', 45, 40, True)  # forraje visible por defecto
            ]

            # Generar malla única para todos (o generar por separado, pero compartir malla ahorra tiempo)
            puntos_malla = self._generar_malla_puntos(gdf_area, densidad=1000)
            if not puntos_malla:
                return None

            for var, nombre, radius, blur, default_show in variables:
                puntos_muestra = resultados.get(f'puntos_{var}', [])
                if not puntos_muestra:
                    continue
                puntos_interpolados = self._interpolar_valores_knn(puntos_muestra, puntos_malla.copy(), var)
                heat_data = []
                for p in puntos_interpolados:
                    if var == 'carbono':
                        val = p.get('carbono_ton_ha', 0)
                    elif var == 'ndvi':
                        val = p.get('ndvi', 0)
                    elif var == 'ndwi':
                        val = p.get('ndwi', 0)
                    elif var == 'biodiversidad':
                        val = p.get('indice_shannon', 0)
                    elif var == 'forraje':
                        val = p.get('productividad_kg_ms_ha', 0)
                    else:
                        continue
                    heat_data.append([p['lat'], p['lon'], val])

                gradient = self.estilos['gradientes'].get(var, self.estilos['gradientes']['carbono'])
                HeatMap(
                    heat_data,
                    name=nombre,
                    min_opacity=0.6,
                    radius=radius,
                    blur=blur,
                    gradient=gradient,
                    max_zoom=18,
                    show=default_show
                ).add_to(m)

            folium.LayerControl().add_to(m)
            return m
        except Exception as e:
            st.warning(f"Error al crear mapa combinado: {str(e)}")
            return None

    def crear_mapa_estatico(self, resultados, variable='carbono', gdf_area=None, dpi=150):
        if not resultados or gdf_area is None or gdf_area.empty:
            return None
        puntos_muestra = resultados.get(f'puntos_{variable}', [])
        if not puntos_muestra:
            return None
        puntos_malla = self._generar_malla_puntos(gdf_area, densidad=800)
        if not puntos_malla:
            return None
        puntos_interpolados = self._interpolar_valores_knn(puntos_muestra, puntos_malla, variable)
        lats = [p['lat'] for p in puntos_interpolados]
        lons = [p['lon'] for p in puntos_interpolados]
        if variable == 'carbono':
            valores = [p['carbono_ton_ha'] for p in puntos_interpolados]
            titulo = 'Carbono (ton C/ha)'
            cmap_name = 'carbono'
        elif variable == 'ndvi':
            valores = [p['ndvi'] for p in puntos_interpolados]
            titulo = 'NDVI'
            cmap_name = 'ndvi'
        elif variable == 'ndwi':
            valores = [p['ndwi'] for p in puntos_interpolados]
            titulo = 'NDWI'
            cmap_name = 'ndwi'
        elif variable == 'biodiversidad':
            valores = [p['indice_shannon'] for p in puntos_interpolados]
            titulo = 'Índice de Shannon'
            cmap_name = 'biodiversidad'
        elif variable == 'forraje':
            valores = [p['productividad_kg_ms_ha'] for p in puntos_interpolados]
            titulo = 'Productividad (kg MS/ha)'
            cmap_name = 'forraje'
        elif variable == 'ndre':
            valores = [p['ndre'] for p in puntos_interpolados]
            titulo = 'NDRE'
            cmap_name = 'ndre'
        elif variable == 'msavi':
            valores = [p['msavi'] for p in puntos_interpolados]
            titulo = 'MSAVI'
            cmap_name = 'msavi'
        elif variable == 'evi':
            valores = [p['evi'] for p in puntos_interpolados]
            titulo = 'EVI'
            cmap_name = 'evi'
        else:
            return None
        bounds = gdf_area.total_bounds
        minx, miny, maxx, maxy = bounds
        grid_x, grid_y = np.mgrid[minx:maxx:100j, miny:maxy:100j]
        grid_z = griddata((lons, lats), valores, (grid_x, grid_y), method='cubic')
        fig, ax = plt.subplots(1, 1, figsize=(10, 8))
        colormap = LinearSegmentedColormap.from_list(cmap_name, list(self.estilos['gradientes'][cmap_name].values()))
        im = ax.imshow(grid_z.T, extent=[minx, maxx, miny, maxy], origin='lower', cmap=colormap, aspect='auto')
        plt.colorbar(im, ax=ax, label=titulo)
        ax.set_title(f'Mapa de {titulo}')
        ax.set_xlabel('Longitud')
        ax.set_ylabel('Latitud')
        ax.grid(True, linestyle='--', alpha=0.5)
        if gdf_area is not None and not gdf_area.empty:
            boundary_geom = gdf_area.geometry.iloc[0].boundary
            if boundary_geom and not boundary_geom.is_empty:
                gpd.GeoSeries([boundary_geom]).plot(ax=ax, color='black', linewidth=1.5)
        buf = BytesIO()
        plt.savefig(buf, format='png', dpi=dpi, bbox_inches='tight')
        plt.close(fig)
        buf.seek(0)
        return buf

# ===============================
# 📊 VISUALIZACIONES
# ===============================
class Visualizaciones:
    _TEMPLATE = {
        'paper_bgcolor': 'rgba(0,0,0,0)',
        'plot_bgcolor': 'rgba(0,0,0,0)',
        'font': {'color': '#cbd5e1', 'family': 'Inter, sans-serif', 'size': 12},
        'title': {'x': 0.5, 'font': {'size': 15, 'color': '#e2e8f0', 'weight': 600}},
        'hovermode': 'x unified',
        'margin': {'t': 40, 'b': 40, 'l': 50, 'r': 20},
        'xaxis': {'gridcolor': 'rgba(255,255,255,0.05)', 'zerolinecolor': 'rgba(255,255,255,0.08)'},
        'yaxis': {'gridcolor': 'rgba(255,255,255,0.05)', 'zerolinecolor': 'rgba(255,255,255,0.08)'},
    }

    @staticmethod
    def _aplicar_estilo(fig, titulo='', height=380):
        fig.update_layout(
            title=dict(text=titulo, x=0.5, font=dict(size=15, color='#e2e8f0', weight=600)),
            paper_bgcolor='rgba(0,0,0,0)',
            plot_bgcolor='rgba(0,0,0,0)',
            font=dict(color='#cbd5e1', family='Inter, sans-serif', size=11),
            hovermode='x unified',
            margin=dict(t=45, b=40, l=50, r=20),
            height=height,
            legend=dict(font=dict(color='#94a3b8', size=10), orientation='h', y=1.12),
        )
        fig.update_xaxes(gridcolor='rgba(255,255,255,0.05)', zerolinecolor='rgba(255,255,255,0.08)',
                         tickfont=dict(color='#94a3b8', size=10))
        fig.update_yaxes(gridcolor='rgba(255,255,255,0.05)', zerolinecolor='rgba(255,255,255,0.08)',
                         tickfont=dict(color='#94a3b8', size=10))
        return fig

    @staticmethod
    def crear_grafico_barras_carbono(desglose: Dict):
        if not desglose:
            fig = go.Figure()
            fig.update_layout(title='No hay datos de carbono disponibles', height=400)
            return fig
        descripciones = {'AGB': 'Biomasa Aérea Viva', 'BGB': 'Biomasa de Raíces', 'DW': 'Madera Muerta', 'LI': 'Hojarasca', 'SOC': 'Carbono Orgánico del Suelo'}
        colores = ['#065f46', '#059669', '#10b981', '#34d399', '#6ee7b7']
        etiquetas = [f"{descripciones.get(k, k)}<br><span style='font-size:0.75rem;color:#94a3b8'>({k})</span>" for k in desglose.keys()]
        fig = go.Figure(data=[go.Bar(
            x=list(desglose.keys()),
            y=list(desglose.values()),
            marker_color=colores[:len(desglose)],
            marker_line=dict(width=0),
            text=[f"{v:.1f}" for v in desglose.values()],
            textposition='outside',
            textfont=dict(color='#e2e8f0', size=12, weight=600),
            hovertemplate='<b>%{x}</b><br>%{y:.1f} ton C/ha<extra></extra>',
        )])
        fig = Visualizaciones._aplicar_estilo(fig, 'Distribución de Carbono por Pools')
        fig.update_yaxes(title='ton C/ha')
        return fig

    @staticmethod
    def crear_grafico_radar_biodiversidad(shannon_data: Dict):
        if not shannon_data:
            fig = go.Figure()
            fig.update_layout(title='No hay datos de biodiversidad disponibles', height=400)
            return fig
        categorias = ['Shannon', 'Riqueza', 'Abundancia', 'Equitatividad', 'Conservación']
        try:
            shannon_norm = min(shannon_data.get('indice_shannon', 0) / 4.0 * 100, 100)
            riqueza_norm = min(shannon_data.get('riqueza_especies', 0) / 200 * 100, 100)
            abundancia_norm = min(shannon_data.get('abundancia_total', 0) / 2000 * 100, 100)
            equitatividad = random.uniform(70, 90)
            conservacion = random.uniform(60, 95)
            valores = [shannon_norm, riqueza_norm, abundancia_norm, equitatividad, conservacion]
            fig = go.Figure(data=go.Scatterpolar(
                r=valores, theta=categorias,
                fill='toself',
                fillcolor='rgba(139, 92, 246, 0.25)',
                line=dict(color='#8b5cf6', width=2),
                marker=dict(size=6, color='#a78bfa'),
                name='Biodiversidad',
            ))
            fig.update_layout(
                paper_bgcolor='rgba(0,0,0,0)', plot_bgcolor='rgba(0,0,0,0)',
                font=dict(color='#cbd5e1', family='Inter, sans-serif', size=11),
                polar=dict(
                    bgcolor='rgba(0,0,0,0)',
                    radialaxis=dict(visible=True, range=[0, 100],
                                    gridcolor='rgba(255,255,255,0.08)',
                                    linecolor='rgba(255,255,255,0.05)',
                                    tickfont=dict(color='#94a3b8', size=9)),
                    angularaxis=dict(gridcolor='rgba(255,255,255,0.05)',
                                     tickfont=dict(color='#94a3b8', size=10)),
                ),
                showlegend=False,
                height=380,
                margin=dict(t=20, b=20, l=40, r=40),
            )
            return fig
        except Exception as e:
            fig = go.Figure()
            fig.update_layout(title='Error al generar gráfico de biodiversidad', height=400)
            return fig

    @staticmethod
    def crear_grafico_comparativo(puntos_carbono, puntos_ndvi, puntos_ndwi, puntos_biodiversidad):
        if not puntos_carbono or not puntos_ndvi:
            return None
        try:
            n = min(50, len(puntos_carbono))
            fig = make_subplots(rows=2, cols=2,
                subplot_titles=('Carbono vs NDVI', 'Carbono vs NDWI', 'Shannon vs NDVI', 'Shannon vs NDWI'),
                vertical_spacing=0.18, horizontal_spacing=0.18)
            carbono_vals = [p['carbono_ton_ha'] for p in puntos_carbono[:n]]
            ndvi_vals = [p['ndvi'] for p in puntos_ndvi[:n]]
            ndwi_vals = [p['ndwi'] for p in puntos_ndwi[:n]]
            shannon_vals = [p['indice_shannon'] for p in puntos_biodiversidad[:n]]
            fig.add_trace(go.Scatter(x=ndvi_vals, y=carbono_vals, mode='markers',
                marker=dict(color='#10b981', size=7, line=dict(width=0.5, color='rgba(255,255,255,0.2)')),
                name='Carbono-NDVI'), row=1, col=1)
            fig.add_trace(go.Scatter(x=ndwi_vals, y=carbono_vals, mode='markers',
                marker=dict(color='#3b82f6', size=7, line=dict(width=0.5, color='rgba(255,255,255,0.2)')),
                name='Carbono-NDWI'), row=1, col=2)
            fig.add_trace(go.Scatter(x=ndvi_vals, y=shannon_vals, mode='markers',
                marker=dict(color='#8b5cf6', size=7, line=dict(width=0.5, color='rgba(255,255,255,0.2)')),
                name='Shannon-NDVI'), row=2, col=1)
            fig.add_trace(go.Scatter(x=ndwi_vals, y=shannon_vals, mode='markers',
                marker=dict(color='#f59e0b', size=7, line=dict(width=0.5, color='rgba(255,255,255,0.2)')),
                name='Shannon-NDWI'), row=2, col=2)
            fig.update_layout(
                paper_bgcolor='rgba(0,0,0,0)', plot_bgcolor='rgba(0,0,0,0)',
                font=dict(color='#cbd5e1', family='Inter, sans-serif', size=11),
                height=680, showlegend=True,
                title=dict(text='Comparación de Variables Ambientales', x=0.5,
                          font=dict(size=15, color='#e2e8f0', weight=600)),
                legend=dict(font=dict(color='#94a3b8', size=9), orientation='h', y=1.02),
            )
            for i in range(1, 3):
                for j in range(1, 3):
                    fig.update_xaxes(gridcolor='rgba(255,255,255,0.05)', row=i, col=j,
                                     tickfont=dict(color='#94a3b8', size=9))
                    fig.update_yaxes(gridcolor='rgba(255,255,255,0.05)', row=i, col=j,
                                     tickfont=dict(color='#94a3b8', size=9))
            return fig
        except Exception as e:
            return None

    @staticmethod
    def crear_grafico_forrajero(disponibilidad_forrajera: Dict, equivalentes_vaca: Dict):
        fig = make_subplots(rows=2, cols=2,
            subplot_titles=('Disponibilidad Forrajera', 'Equivalentes Vaca',
                           'Distribución por Sublote', 'Plan de Rotación'),
            specs=[[{'type': 'bar'}, {'type': 'pie'}],
                   [{'type': 'bar'}, {'type': 'table'}]],
            vertical_spacing=0.18, horizontal_spacing=0.18, row_heights=[0.5, 0.5])
        fig.add_trace(go.Bar(
            x=['Productividad', 'Disponible Total', 'Aprovechable'],
            y=[disponibilidad_forrajera.get('productividad_kg_ms_ha', 0),
               disponibilidad_forrajera.get('disponibilidad_total_kg_ms', 0) / 1000,
               disponibilidad_forrajera.get('forraje_aprovechable_kg_ms', 0) / 1000],
            marker_color=['#92400e', '#d97706', '#f59e0b'],
            marker_line=dict(width=0),
            texttemplate='%{y:,.0f}', textposition='outside',
            textfont=dict(color='#e2e8f0', size=10),
            hovertemplate='<b>%{x}</b><br>%{y:,.1f}<extra></extra>',
        ), row=1, col=1)
        fig.add_trace(go.Pie(
            labels=['EV por día', 'EV para período', 'EV recomendado'],
            values=[equivalentes_vaca.get('ev_por_dia', 0),
                    equivalentes_vaca.get('ev_para_periodo', 0),
                    equivalentes_vaca.get('ev_recomendado', 0)],
            marker=dict(colors=['#f97316', '#a855f7', '#06b6d4']),
            textinfo='label+value', textfont=dict(color='#f1f5f9', size=10),
            hole=0.45, hovertemplate='%{label}: %{value:.1f}<extra></extra>',
        ), row=1, col=2)
        fig.update_layout(
            paper_bgcolor='rgba(0,0,0,0)', plot_bgcolor='rgba(0,0,0,0)',
            font=dict(color='#cbd5e1', family='Inter, sans-serif', size=11),
            height=680,
            title=dict(text='Análisis Forrajero Completo', x=0.5,
                      font=dict(size=15, color='#e2e8f0', weight=600)),
            legend=dict(font=dict(color='#94a3b8', size=9), orientation='h', y=1.02),
            showlegend=True,
        )
        fig.update_xaxes(gridcolor='rgba(255,255,255,0.05)', tickfont=dict(color='#94a3b8', size=9))
        fig.update_yaxes(gridcolor='rgba(255,255,255,0.05)', tickfont=dict(color='#94a3b8', size=9),
                         title='kg MS/ha / ton MS')
        return fig

    @staticmethod
    def crear_metricas_kpi(carbono_total: float, co2_total: float, shannon: float, area: float,
                            ndvi: float = 0, ndwi: float = 0, forraje_kg: float = 0, ev: float = 0):
        def _card(icon, label, value, unit, gradient, badge=None):
            badge_html = f'<span style="background:rgba(255,255,255,0.15);padding:2px 10px;border-radius:20px;font-size:0.7rem;margin-left:6px;">{badge}</span>' if badge else ''
            return (
                '<div style="background:' + gradient + ';padding:1.5rem 1.75rem;border-radius:20px;'
                'border:1px solid rgba(255,255,255,0.1);backdrop-filter:blur(8px);'
                'box-shadow:0 8px 30px rgba(0,0,0,0.3);transition:all 0.25s ease;">'
                '<div style="display:flex;align-items:center;gap:0.5rem;margin-bottom:0.5rem;">'
                '<span style="font-size:1.3rem;">' + icon + '</span>'
                '<span style="color:rgba(255,255,255,0.7);font-size:0.75rem;font-weight:500;text-transform:uppercase;letter-spacing:0.05em;">'
                + label + badge_html + '</span></div>'
                '<div style="font-size:2.5rem;font-weight:800;color:white;letter-spacing:-0.03em;line-height:1.1;">'
                + value + '</div>'
                '<div style="color:rgba(255,255,255,0.5);font-size:0.75rem;margin-top:0.25rem;">' + unit + '</div></div>'
            )

        cards = []
        cards.append(_card("🌳", "Carbono Total", f"{carbono_total:,.0f}", "ton C almacenadas",
                           "linear-gradient(135deg, #065f46 0%, #059669 100%)"))
        cards.append(_card("🏭", "CO₂ Equivalente", f"{co2_total:,.0f}", "ton CO₂e — potencial créditos",
                           "linear-gradient(135deg, #0a7e5a 0%, #10b981 100%)"))
        cards.append(_card("🦋", "Biodiversidad", f"{shannon:.2f}", "Índice de Shannon",
                           "linear-gradient(135deg, #6d28d9 0%, #8b5cf6 100%)",
                           "bajo" if shannon < 1.5 else ("medio" if shannon < 2.5 else "alto")))
        cards.append(_card("📐", "Área de Estudio", f"{area:,.1f}", "hectáreas",
                           "linear-gradient(135deg, #1e40af 0%, #3b82f6 100%)"))

        if forraje_kg > 0:
            cards.append(_card("🌿", "Productividad Forrajera", f"{forraje_kg:,.0f}", "kg MS/ha",
                               "linear-gradient(135deg, #92400e 0%, #d97706 100%)"))
        if ev > 0:
            cards.append(_card("🐄", "Carga Animal Recomendada", f"{ev:.1f}", "EV para 30 días",
                               "linear-gradient(135deg, #831843 0%, #db2777 100%)"))
        cards.append(_card("📈", "NDVI Promedio", f"{ndvi:.3f}", "Salud de la vegetación",
                           "linear-gradient(135deg, #14532d 0%, #16a34a 100%)"))
        cards.append(_card("💧", "NDWI Promedio", f"{ndwi:.3f}", "Contenido de agua",
                           "linear-gradient(135deg, #0c4a6e 0%, #0284c7 100%)"))

        return (
            '<div style="display:grid;grid-template-columns:repeat(auto-fit,minmax(220px,1fr));gap:1rem;margin-bottom:1.5rem;">'
            + ''.join(cards) + '</div>'
        )

# ===============================
# 📄 GENERADOR DE REPORTES
# ===============================
class GeneradorReportes:
    def __init__(self, resultados, gdf, sistema_mapas=None):
        self.resultados = resultados
        self.gdf = gdf
        self.sistema_mapas = sistema_mapas
        self.buffer_pdf = BytesIO()
        self.buffer_docx = BytesIO()

    def _fig_to_png(self, fig, width=800, height=500):
        if fig is None:
            return None
        try:
            img_bytes = fig.to_image(format='png', width=width, height=height, scale=2)
            return BytesIO(img_bytes)
        except Exception as e:
            st.warning(f"No se pudo convertir el gráfico a PNG: {str(e)}")
            return None

    def _mapa_to_png(self, mapa, width=800, height=600):
        try:
            if mapa is None:
                return None
            from PIL import Image, ImageDraw
            img = Image.new('RGB', (width, height), color='white')
            draw = ImageDraw.Draw(img)
            draw.text((width//2 - 100, height//2 - 20), "Mapa interactivo", fill='black')
            draw.text((width//2 - 150, height//2 + 10), "Disponible en la aplicación web", fill='gray')
            draw.rectangle([10, 10, width-10, height-10], outline='blue', width=3)
            img_byte_arr = BytesIO()
            img.save(img_byte_arr, format='PNG')
            img_byte_arr.seek(0)
            return img_byte_arr
        except Exception as e:
            st.warning(f"No se pudo convertir el mapa a PNG: {str(e)}")
            return None

    def _crear_graficos(self):
        vis = Visualizaciones()
        res = self.resultados
        graficos = {}
        if 'desglose_promedio' in res and res['desglose_promedio']:
            fig_carbono = vis.crear_grafico_barras_carbono(res['desglose_promedio'])
            graficos['carbono'] = self._fig_to_png(fig_carbono)
        if 'puntos_biodiversidad' in res and res['puntos_biodiversidad'] and len(res['puntos_biodiversidad']) > 0:
            fig_biodiv = vis.crear_grafico_radar_biodiversidad(res['puntos_biodiversidad'][0])
            graficos['biodiv'] = self._fig_to_png(fig_biodiv)
        if all(k in res for k in ['puntos_carbono', 'puntos_ndvi', 'puntos_ndwi', 'puntos_biodiversidad']):
            fig_comparativo = vis.crear_grafico_comparativo(
                res['puntos_carbono'], res['puntos_ndvi'], res['puntos_ndwi'], res['puntos_biodiversidad']
            )
            if fig_comparativo:
                graficos['comparativo'] = self._fig_to_png(fig_comparativo)
        if 'analisis_forrajero' in res:
            forrajero_data = res['analisis_forrajero']
            if 'disponibilidad_forrajera' in forrajero_data and 'equivalentes_vaca' in forrajero_data:
                fig_forrajero = vis.crear_grafico_forrajero(
                    forrajero_data['disponibilidad_forrajera'],
                    forrajero_data['equivalentes_vaca']
                )
                graficos['forrajero'] = self._fig_to_png(fig_forrajero)
        return graficos

    def generar_pdf(self):
        if not REPORTPDF_AVAILABLE:
            st.error("ReportLab no está instalado. No se puede generar PDF.")
            return None
        try:
            doc = SimpleDocTemplate(self.buffer_pdf, pagesize=A4, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=72)
            story = []
            styles = getSampleStyleSheet()
            titulo_style = ParagraphStyle('TituloPrincipal', parent=styles['Heading1'], fontSize=24, textColor=colors.HexColor('#0a7e5a'), spaceAfter=30, alignment=TA_CENTER)
            subtitulo_style = ParagraphStyle('Subtitulo', parent=styles['Heading2'], fontSize=18, textColor=colors.HexColor('#065f46'), spaceAfter=12, spaceBefore=20)
            seccion_style = ParagraphStyle('Seccion', parent=styles['Heading3'], fontSize=14, textColor=colors.HexColor('#1d4ed8'), spaceAfter=10, spaceBefore=15)
            # Portada
            story.append(Paragraph("INFORME AMBIENTAL INTEGRAL", titulo_style))
            story.append(Spacer(1, 12))
            story.append(Paragraph("Sistema Satelital de Análisis Ambiental", styles['Title']))
            story.append(Spacer(1, 6))
            story.append(Paragraph("Carbono + Biodiversidad + Análisis Forrajero", styles['Heading2']))
            story.append(Spacer(1, 24))
            story.append(Paragraph(f"Fecha de generación: {datetime.now().strftime('%d/%m/%Y %H:%M')}", styles['Normal']))
            story.append(Spacer(1, 36))
            # Resumen ejecutivo
            story.append(Paragraph("RESUMEN EJECUTIVO", subtitulo_style))
            res = self.resultados
            datos_resumen = [
                ["Métrica", "Valor", "Interpretación"],
                ["Área total", f"{res.get('area_total_ha', 0):,.1f} ha", "Superficie del área de estudio"],
                ["Carbono total almacenado", f"{res.get('carbono_total_ton', 0):,.0f} ton C", "Carbono almacenado en el área"],
                ["CO₂ equivalente", f"{res.get('co2_total_ton', 0):,.0f} ton CO₂e", "Potencial de créditos de carbono"],
                ["Índice de Shannon promedio", f"{res.get('shannon_promedio', 0):.3f}", "Nivel de biodiversidad"],
                ["NDVI promedio", f"{res.get('ndvi_promedio', 0):.3f}", "Salud de la vegetación"],
                ["NDWI promedio", f"{res.get('ndwi_promedio', 0):.3f}", "Contenido de agua"],
                ["Tipo de ecosistema", res.get('tipo_ecosistema', 'N/A'), "Ecosistema predominante"],
                ["Puntos de muestreo", str(res.get('num_puntos', 0)), "Muestras analizadas"]
            ]
            tabla_resumen = Table(datos_resumen, colWidths=[150, 120, 200])
            tabla_resumen.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#065f46')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 11),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#f0f9ff')),
                ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#cbd5e1')),
                ('ALIGN', (0, 1), (-1, -1), 'LEFT'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ]))
            story.append(tabla_resumen)
            story.append(Spacer(1, 20))
            # Análisis de carbono
            story.append(PageBreak())
            story.append(Paragraph("ANÁLISIS DE CARBONO", subtitulo_style))
            if res.get('desglose_promedio'):
                descripciones = {
                    'AGB': 'Biomasa Aérea Viva', 'BGB': 'Biomasa de Raíces', 'DW': 'Madera Muerta',
                    'LI': 'Hojarasca', 'SOC': 'Carbono Orgánico del Suelo'
                }
                datos_carbono = [["Pool", "Descripción", "Ton C/ha", "Porcentaje"]]
                total = sum(res['desglose_promedio'].values())
                for pool, valor in res['desglose_promedio'].items():
                    porcentaje = (valor / total * 100) if total > 0 else 0
                    datos_carbono.append([pool, descripciones.get(pool, pool), f"{valor:.2f}", f"{porcentaje:.1f}%"])
                tabla_carbono = Table(datos_carbono, colWidths=[60, 180, 70, 70])
                tabla_carbono.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#0a7e5a')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                    ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                    ('ALIGN', (2, 1), (3, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#f0fdf4')),
                    ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#d1fae5')),
                ]))
                story.append(tabla_carbono)
                story.append(Spacer(1, 15))
            # Análisis de biodiversidad
            story.append(PageBreak())
            story.append(Paragraph("ANÁLISIS DE BIODIVERSIDAD", subtitulo_style))
            if res.get('puntos_biodiversidad') and len(res['puntos_biodiversidad']) > 0:
                biodiv = res['puntos_biodiversidad'][0]
                datos_biodiv = [
                    ["Métrica", "Valor", "Interpretación"],
                    ["Índice de Shannon", f"{biodiv.get('indice_shannon', 0):.3f}", biodiv.get('categoria', 'N/A')],
                    ["Riqueza de especies", str(biodiv.get('riqueza_especies', 0)), "Número estimado de especies"],
                    ["Abundancia total", f"{biodiv.get('abundancia_total', 0):,}", "Individuos estimados"],
                    ["Categoría", biodiv.get('categoria', 'N/A'), "Clasificación según Shannon"]
                ]
                tabla_biodiv = Table(datos_biodiv, colWidths=[120, 100, 180])
                tabla_biodiv.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#8b5cf6')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                    ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                    ('ALIGN', (1, 1), (1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#faf5ff')),
                    ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#e9d5ff')),
                ]))
                story.append(tabla_biodiv)
                story.append(Spacer(1, 15))
            # Análisis forrajero
            story.append(PageBreak())
            story.append(Paragraph("ANÁLISIS FORRAJERO", subtitulo_style))
            if 'analisis_forrajero' in res:
                forrajero_data = res['analisis_forrajero']
                if 'disponibilidad_forrajera' in forrajero_data:
                    disp = forrajero_data['disponibilidad_forrajera']
                    datos_forraje = [
                        ["Métrica", "Valor", "Unidad"],
                        ["Productividad", f"{disp.get('productividad_kg_ms_ha', 0):,.0f}", "kg MS/ha"],
                        ["Disponibilidad total", f"{disp.get('disponibilidad_total_kg_ms', 0)/1000:,.1f}", "ton MS"],
                        ["Forraje aprovechable", f"{disp.get('forraje_aprovechable_kg_ms', 0)/1000:,.1f}", "ton MS"],
                        ["Tasa crecimiento diario", f"{disp.get('tasa_crecimiento_diario_kg', 0):,.0f}", "kg/día"],
                        ["Categoría productividad", disp.get('categoria_productividad', 'N/A').title(), ""]
                    ]
                    tabla_forraje = Table(datos_forraje, colWidths=[150, 100, 80])
                    tabla_forraje.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#8B4513')),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                        ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                        ('ALIGN', (1, 1), (2, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 10),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#fdf4e3')),
                        ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#d2b48c')),
                    ]))
                    story.append(tabla_forraje)
                    story.append(Spacer(1, 15))
                if 'equivalentes_vaca' in forrajero_data:
                    ev = forrajero_data['equivalentes_vaca']
                    datos_ev = [
                        ["Concepto", "Valor"],
                        ["EV por día", f"{ev.get('ev_por_dia', 0):.1f}"],
                        ["EV para 30 días", f"{ev.get('ev_para_periodo', 0):.1f}"],
                        ["EV recomendado", f"{ev.get('ev_recomendado', 0):.1f}"],
                        ["Consumo EV diario", f"{ev.get('consumo_ev_diario_kg', 0)} kg"]
                    ]
                    tabla_ev = Table(datos_ev, colWidths=[150, 100])
                    tabla_ev.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#CD853F')),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 10),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#fff8dc')),
                        ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#cd853f')),
                    ]))
                    story.append(tabla_ev)
                    story.append(Spacer(1, 15))
                if 'sublotes' in forrajero_data and forrajero_data['sublotes']:
                    datos_sublotes = [["Sublote", "Área (ha)", "Productividad (kg MS/ha)", "Forraje aprovechable (ton)"]]
                    for s in forrajero_data['sublotes']:
                        datos_sublotes.append([
                            str(s['sublote_id']),
                            f"{s['area_ha']:.1f}",
                            f"{s['disponibilidad_kg_ms_ha']:,.0f}",
                            f"{s['forraje_aprovechable_kg_ms']/1000:.1f}"
                        ])
                    tabla_sublotes = Table(datos_sublotes, colWidths=[60, 70, 120, 100])
                    tabla_sublotes.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#8B4513')),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                        ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                        ('ALIGN', (1, 1), (3, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 10),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#fdf4e3')),
                        ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#d2b48c')),
                    ]))
                    story.append(tabla_sublotes)
                    story.append(Spacer(1, 15))
            # Índices espectrales
            story.append(PageBreak())
            story.append(Paragraph("ÍNDICES ESPECTRALES", subtitulo_style))
            datos_indices = [
                ["Índice", "Valor promedio"],
                ["NDVI", f"{res.get('ndvi_promedio', 0):.3f}"],
                ["NDWI", f"{res.get('ndwi_promedio', 0):.3f}"],
            ]
            if 'puntos_ndre' in res:
                ndre_vals = [p['ndre'] for p in res['puntos_ndre']]
                datos_indices.append(["NDRE", f"{np.mean(ndre_vals):.3f}"])
            if 'puntos_msavi' in res:
                msavi_vals = [p['msavi'] for p in res['puntos_msavi']]
                datos_indices.append(["MSAVI", f"{np.mean(msavi_vals):.3f}"])
            if 'puntos_evi' in res:
                evi_vals = [p['evi'] for p in res['puntos_evi']]
                datos_indices.append(["EVI", f"{np.mean(evi_vals):.3f}"])
            tabla_indices = Table(datos_indices, colWidths=[100, 100])
            tabla_indices.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#10b981')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#f0fdf4')),
                ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#bbf7d0')),
            ]))
            story.append(tabla_indices)
            story.append(Spacer(1, 20))
            # Mapas estáticos (si los hay)
            if self.sistema_mapas:
                story.append(PageBreak())
                story.append(Paragraph("MAPAS DE CALOR", subtitulo_style))
                variables = ['carbono', 'ndvi', 'ndwi', 'biodiversidad', 'forraje']
                for var in variables:
                    mapa = self.sistema_mapas.crear_mapa_estatico(self.resultados, var, self.gdf)
                    if mapa:
                        story.append(Paragraph(f"Mapa de {var.replace('_',' ').title()}", seccion_style))
                        story.append(Image(mapa, width=450, height=350))
                        story.append(Spacer(1, 12))
            # Conclusiones
            story.append(PageBreak())
            story.append(Paragraph("CONCLUSIONES Y RECOMENDACIONES", subtitulo_style))
            if 'analisis_forrajero' in res:
                forrajero_data = res['analisis_forrajero']
            else:
                forrajero_data = {}
            conclusiones = [
                f"El área de estudio de {res.get('area_total_ha', 0):,.1f} hectáreas almacena {res.get('carbono_total_ton', 0):,.0f} ton C, equivalente a {res.get('co2_total_ton', 0):,.0f} ton CO₂e.",
                f"El índice de Shannon promedio es {res.get('shannon_promedio', 0):.3f}, lo que indica una biodiversidad {res.get('puntos_biodiversidad', [{}])[0].get('categoria', 'N/A').lower()}.",
                f"El NDVI promedio de {res.get('ndvi_promedio', 0):.3f} sugiere una cobertura vegetal moderada.",
                f"La productividad forrajera estimada es de {forrajero_data.get('disponibilidad_forrajera', {}).get('productividad_kg_ms_ha', 0):,.0f} kg MS/ha, lo que permite recomendar una carga de {forrajero_data.get('equivalentes_vaca', {}).get('ev_recomendado', 0):.1f} EV para un período de 30 días."
            ]
            for conc in conclusiones:
                story.append(Paragraph(conc, styles['Normal']))
                story.append(Spacer(1, 8))
            doc.build(story)
            self.buffer_pdf.seek(0)
            return self.buffer_pdf
        except Exception as e:
            st.error(f"Error generando PDF: {str(e)}")
            import traceback
            st.error(traceback.format_exc())
            return None

    def generar_docx(self):
        if not REPORTDOCX_AVAILABLE:
            st.error("python-docx no está instalado. No se puede generar DOCX.")
            return None
        try:
            doc = Document()
            style = doc.styles['Normal']
            style.font.name = 'Arial'
            style.font.size = Pt(11)
            title = doc.add_heading('INFORME AMBIENTAL INTEGRAL', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph(f"Fecha de generación: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
            doc.add_paragraph()
            # Resumen ejecutivo
            doc.add_heading('RESUMEN EJECUTIVO', level=1)
            res = self.resultados
            tabla_resumen = doc.add_table(rows=9, cols=3)
            tabla_resumen.style = 'Light Shading'
            tabla_resumen.cell(0, 0).text = 'Métrica'
            tabla_resumen.cell(0, 1).text = 'Valor'
            tabla_resumen.cell(0, 2).text = 'Interpretación'
            datos = [
                ('Área total', f"{res.get('area_total_ha', 0):,.1f} ha", 'Superficie del área de estudio'),
                ('Carbono total almacenado', f"{res.get('carbono_total_ton', 0):,.0f} ton C", 'Carbono almacenado en el área'),
                ('CO₂ equivalente', f"{res.get('co2_total_ton', 0):,.0f} ton CO₂e", 'Potencial de créditos de carbono'),
                ('Índice de Shannon promedio', f"{res.get('shannon_promedio', 0):.3f}", 'Nivel de biodiversidad'),
                ('NDVI promedio', f"{res.get('ndvi_promedio', 0):.3f}", 'Salud de la vegetación'),
                ('NDWI promedio', f"{res.get('ndwi_promedio', 0):.3f}", 'Contenido de agua'),
                ('Tipo de ecosistema', res.get('tipo_ecosistema', 'N/A'), 'Ecosistema predominante'),
                ('Puntos de muestreo', str(res.get('num_puntos', 0)), 'Muestras analizadas')
            ]
            for i, (met, val, interp) in enumerate(datos, 1):
                tabla_resumen.cell(i, 0).text = met
                tabla_resumen.cell(i, 1).text = val
                tabla_resumen.cell(i, 2).text = interp
            doc.add_paragraph()
            # Análisis de carbono
            doc.add_heading('ANÁLISIS DE CARBONO', level=1)
            if res.get('desglose_promedio'):
                doc.add_heading('Distribución por Pools', level=2)
                tabla_carbono = doc.add_table(rows=6, cols=4)
                tabla_carbono.style = 'Light Shading'
                tabla_carbono.cell(0, 0).text = 'Pool'
                tabla_carbono.cell(0, 1).text = 'Descripción'
                tabla_carbono.cell(0, 2).text = 'Ton C/ha'
                tabla_carbono.cell(0, 3).text = 'Porcentaje'
                desc = {'AGB': 'Biomasa Aérea Viva', 'BGB': 'Biomasa de Raíces', 'DW': 'Madera Muerta', 'LI': 'Hojarasca', 'SOC': 'Carbono Orgánico del Suelo'}
                total = sum(res['desglose_promedio'].values())
                for i, (pool, valor) in enumerate(res['desglose_promedio'].items(), 1):
                    tabla_carbono.cell(i, 0).text = pool
                    tabla_carbono.cell(i, 1).text = desc.get(pool, pool)
                    tabla_carbono.cell(i, 2).text = f"{valor:.2f}"
                    porcentaje = (valor / total * 100) if total > 0 else 0
                    tabla_carbono.cell(i, 3).text = f"{porcentaje:.1f}%"
            doc.add_page_break()
            # Análisis de biodiversidad
            doc.add_heading('ANÁLISIS DE BIODIVERSIDAD', level=1)
            if res.get('puntos_biodiversidad') and len(res['puntos_biodiversidad']) > 0:
                biodiv = res['puntos_biodiversidad'][0]
                tabla_biodiv = doc.add_table(rows=5, cols=3)
                tabla_biodiv.style = 'Light Shading'
                tabla_biodiv.cell(0, 0).text = 'Métrica'
                tabla_biodiv.cell(0, 1).text = 'Valor'
                tabla_biodiv.cell(0, 2).text = 'Interpretación'
                datos_biodiv = [
                    ('Índice de Shannon', f"{biodiv.get('indice_shannon', 0):.3f}", biodiv.get('categoria', 'N/A')),
                    ('Riqueza de especies', str(biodiv.get('riqueza_especies', 0)), 'Número estimado de especies'),
                    ('Abundancia total', f"{biodiv.get('abundancia_total', 0):,}", 'Individuos estimados'),
                    ('Categoría', biodiv.get('categoria', 'N/A'), 'Clasificación según Shannon')
                ]
                for i, (met, val, interp) in enumerate(datos_biodiv, 1):
                    tabla_biodiv.cell(i, 0).text = met
                    tabla_biodiv.cell(i, 1).text = val
                    tabla_biodiv.cell(i, 2).text = interp
            doc.add_page_break()
            # Análisis forrajero
            doc.add_heading('ANÁLISIS FORRAJERO', level=1)
            if 'analisis_forrajero' in res:
                forrajero_data = res['analisis_forrajero']
                if 'disponibilidad_forrajera' in forrajero_data:
                    disp = forrajero_data['disponibilidad_forrajera']
                    doc.add_heading('Disponibilidad Forrajera', level=2)
                    tabla_forraje = doc.add_table(rows=6, cols=3)
                    tabla_forraje.style = 'Light Shading'
                    tabla_forraje.cell(0, 0).text = 'Métrica'
                    tabla_forraje.cell(0, 1).text = 'Valor'
                    tabla_forraje.cell(0, 2).text = 'Unidad'
                    datos_f = [
                        ('Productividad', f"{disp.get('productividad_kg_ms_ha', 0):,.0f}", 'kg MS/ha'),
                        ('Disponibilidad total', f"{disp.get('disponibilidad_total_kg_ms', 0)/1000:,.1f}", 'ton MS'),
                        ('Forraje aprovechable', f"{disp.get('forraje_aprovechable_kg_ms', 0)/1000:,.1f}", 'ton MS'),
                        ('Tasa crecimiento diario', f"{disp.get('tasa_crecimiento_diario_kg', 0):,.0f}", 'kg/día'),
                        ('Categoría productividad', disp.get('categoria_productividad', 'N/A').title(), '')
                    ]
                    for i, (met, val, uni) in enumerate(datos_f, 1):
                        tabla_forraje.cell(i, 0).text = met
                        tabla_forraje.cell(i, 1).text = val
                        tabla_forraje.cell(i, 2).text = uni
                if 'equivalentes_vaca' in forrajero_data:
                    ev = forrajero_data['equivalentes_vaca']
                    doc.add_heading('Equivalentes Vaca', level=2)
                    tabla_ev = doc.add_table(rows=5, cols=2)
                    tabla_ev.style = 'Light Shading'
                    tabla_ev.cell(0, 0).text = 'Concepto'
                    tabla_ev.cell(0, 1).text = 'Valor'
                    datos_ev = [
                        ('EV por día', f"{ev.get('ev_por_dia', 0):.1f}"),
                        ('EV para 30 días', f"{ev.get('ev_para_periodo', 0):.1f}"),
                        ('EV recomendado', f"{ev.get('ev_recomendado', 0):.1f}"),
                        ('Consumo EV diario', f"{ev.get('consumo_ev_diario_kg', 0)} kg")
                    ]
                    for i, (concepto, valor) in enumerate(datos_ev, 1):
                        tabla_ev.cell(i, 0).text = concepto
                        tabla_ev.cell(i, 1).text = valor
                if 'sublotes' in forrajero_data and forrajero_data['sublotes']:
                    doc.add_heading('Sublotes', level=2)
                    tabla_sub = doc.add_table(rows=len(forrajero_data['sublotes'])+1, cols=4)
                    tabla_sub.style = 'Light Shading'
                    tabla_sub.cell(0, 0).text = 'Sublote'
                    tabla_sub.cell(0, 1).text = 'Área (ha)'
                    tabla_sub.cell(0, 2).text = 'Productividad (kg MS/ha)'
                    tabla_sub.cell(0, 3).text = 'Forraje aprovechable (ton)'
                    for i, s in enumerate(forrajero_data['sublotes'], 1):
                        tabla_sub.cell(i, 0).text = str(s['sublote_id'])
                        tabla_sub.cell(i, 1).text = f"{s['area_ha']:.1f}"
                        tabla_sub.cell(i, 2).text = f"{s['disponibilidad_kg_ms_ha']:,.0f}"
                        tabla_sub.cell(i, 3).text = f"{s['forraje_aprovechable_kg_ms']/1000:.1f}"
            doc.save(self.buffer_docx)
            self.buffer_docx.seek(0)
            return self.buffer_docx
        except Exception as e:
            st.error(f"Error generando DOCX: {str(e)}")
            return None

    def generar_geojson(self):
        try:
            gdf_out = self.gdf.copy()
            res = self.resultados
            if res:
                gdf_out['area_ha'] = res.get('area_total_ha', 0)
                gdf_out['carbono_total_ton'] = res.get('carbono_total_ton', 0)
                gdf_out['shannon_promedio'] = res.get('shannon_promedio', 0)
                gdf_out['ecosistema'] = res.get('tipo_ecosistema', 'N/A')
                if 'analisis_forrajero' in res:
                    gdf_out['forraje_kg_ms_ha'] = res['analisis_forrajero']['disponibilidad_forrajera']['productividad_kg_ms_ha']
            geojson_str = gdf_out.to_json()
            return geojson_str
        except Exception as e:
            st.error(f"Error generando GeoJSON: {str(e)}")
            return json.dumps({"error": str(e)})

# ===============================
# FUNCIÓN PARA GENERAR INFORME CON IA (ahora usando Groq)
# ===============================
def generar_reporte_ia(resultados, gdf, sistema_mapas=None):
    """
    Genera un informe en Word con análisis de IA usando Groq.
    """
    import tempfile
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from datetime import datetime
    import io
    import os

    if not REPORTDOCX_AVAILABLE:
        st.error("python-docx no está instalado. No se puede generar el informe.")
        return None

    with tempfile.TemporaryDirectory() as tmpdir:
        doc = Document()
        section = doc.sections[0]
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

        title = doc.add_heading('INFORME AMBIENTAL CON ANÁLISIS DE IA (GROQ)', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle = doc.add_paragraph(f'Fecha: {datetime.now().strftime("%d/%m/%Y %H:%M")}')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        df, stats = preparar_resumen(resultados)

        # 1. Resumen ejecutivo
        doc.add_heading('1. RESUMEN EJECUTIVO', level=1)
        tabla_resumen = doc.add_table(rows=1, cols=3)
        tabla_resumen.style = 'Light Shading'
        tabla_resumen.cell(0, 0).text = 'Métrica'
        tabla_resumen.cell(0, 1).text = 'Valor'
        tabla_resumen.cell(0, 2).text = 'Interpretación'

        metricas = [
            ('Área total', f"{stats['area_total_ha']:,.1f} ha", 'Superficie del área de estudio'),
            ('Carbono total', f"{stats['carbono_total_ton']:,.0f} ton C", 'Almacenamiento total de carbono'),
            ('CO₂ equivalente', f"{stats['co2_total_ton']:,.0f} ton CO₂e", 'Potencial de créditos de carbono'),
            ('Índice Shannon', f"{stats['shannon_promedio']:.3f}", 'Nivel de biodiversidad'),
            ('NDVI promedio', f"{stats['ndvi_promedio']:.3f}", 'Salud de la vegetación'),
            ('NDWI promedio', f"{stats['ndwi_promedio']:.3f}", 'Contenido de agua'),
            ('Tipo ecosistema', stats['tipo_ecosistema'], 'Vegetación predominante'),
            ('Puntos muestreo', str(stats['num_puntos']), 'Muestras analizadas')
        ]
        for i, (met, val, interp) in enumerate(metricas, 1):
            row = tabla_resumen.add_row().cells
            row[0].text = met
            row[1].text = val
            row[2].text = interp
        doc.add_paragraph()

        # 2. Análisis de Carbono (usando función de ia_integration)
        doc.add_heading('2. ANÁLISIS DE CARBONO', level=1)
        if resultados.get('desglose_promedio'):
            doc.add_heading('Distribución por pools', level=2)
            tabla_pools = doc.add_table(rows=1, cols=3)
            tabla_pools.style = 'Light Shading'
            tabla_pools.cell(0, 0).text = 'Pool'
            tabla_pools.cell(0, 1).text = 'Descripción'
            tabla_pools.cell(0, 2).text = 'Ton C/ha'
            desc = {'AGB':'Biomasa Aérea Viva', 'BGB':'Biomasa de Raíces', 'DW':'Madera Muerta', 'LI':'Hojarasca', 'SOC':'Carbono Orgánico del Suelo'}
            for pool, valor in resultados['desglose_promedio'].items():
                row = tabla_pools.add_row().cells
                row[0].text = pool
                row[1].text = desc.get(pool, pool)
                row[2].text = f"{valor:.2f}"
            doc.add_paragraph()
            vis = Visualizaciones()
            fig_carbono = vis.crear_grafico_barras_carbono(resultados['desglose_promedio'])
            if fig_carbono:
                try:
                    img_bytes = fig_carbono.to_image(format='png', width=800, height=500, scale=2)
                    img_path = os.path.join(tmpdir, 'carbono.png')
                    with open(img_path, 'wb') as f:
                        f.write(img_bytes)
                    doc.add_picture(img_path, width=Inches(5))
                    doc.add_paragraph()
                except:
                    pass

        doc.add_heading('2.1 Interpretación técnica', level=2)
        analisis_carbono = generar_analisis_carbono(df, stats)
        doc.add_paragraph(analisis_carbono)

        # 3. Análisis de Biodiversidad
        doc.add_heading('3. ANÁLISIS DE BIODIVERSIDAD', level=1)
        if resultados.get('puntos_biodiversidad'):
            biodiv = resultados['puntos_biodiversidad'][0]
            tabla_biodiv = doc.add_table(rows=1, cols=2)
            tabla_biodiv.style = 'Light Shading'
            tabla_biodiv.cell(0, 0).text = 'Métrica'
            tabla_biodiv.cell(0, 1).text = 'Valor'
            metricas_bio = [
                ('Índice Shannon', f"{biodiv.get('indice_shannon', 0):.3f}"),
                ('Categoría', biodiv.get('categoria', 'N/A')),
                ('Riqueza de especies', str(biodiv.get('riqueza_especies', 0))),
                ('Abundancia total', f"{biodiv.get('abundancia_total', 0):,}")
            ]
            for met, val in metricas_bio:
                row = tabla_biodiv.add_row().cells
                row[0].text = met
                row[1].text = val
            doc.add_paragraph()
            fig_biodiv = vis.crear_grafico_radar_biodiversidad(biodiv)
            if fig_biodiv:
                try:
                    img_bytes = fig_biodiv.to_image(format='png', width=800, height=800, scale=2)
                    img_path = os.path.join(tmpdir, 'biodiv.png')
                    with open(img_path, 'wb') as f:
                        f.write(img_bytes)
                    doc.add_picture(img_path, width=Inches(5))
                    doc.add_paragraph()
                except:
                    pass

        doc.add_heading('3.1 Interpretación técnica', level=2)
        analisis_biodiv = generar_analisis_biodiversidad(df, stats)
        doc.add_paragraph(analisis_biodiv)

        # 4. Análisis de Índices Espectrales
        doc.add_heading('4. ANÁLISIS DE ÍNDICES ESPECTRALES', level=1)
        doc.add_heading('4.1 Interpretación técnica', level=2)
        analisis_espectral = generar_analisis_espectral(df, stats)
        doc.add_paragraph(analisis_espectral)

        # 5. Análisis Forrajero
        doc.add_heading('5. ANÁLISIS FORRAJERO', level=1)
        if 'analisis_forrajero' in resultados:
            forrajero = resultados['analisis_forrajero']
            disp = forrajero['disponibilidad_forrajera']
            ev = forrajero['equivalentes_vaca']
            tabla_forraje = doc.add_table(rows=1, cols=2)
            tabla_forraje.style = 'Light Shading'
            tabla_forraje.cell(0, 0).text = 'Métrica'
            tabla_forraje.cell(0, 1).text = 'Valor'
            datos_f = [
                ('Productividad (kg MS/ha)', f"{disp['productividad_kg_ms_ha']:,.0f}"),
                ('Forraje aprovechable (ton)', f"{disp['forraje_aprovechable_kg_ms']/1000:.1f}"),
                ('EV por día', f"{ev['ev_por_dia']:.1f}"),
                ('EV recomendado (30 días)', f"{ev['ev_recomendado']:.1f}")
            ]
            for met, val in datos_f:
                row = tabla_forraje.add_row().cells
                row[0].text = met
                row[1].text = val
            doc.add_paragraph()
            fig_forrajero = vis.crear_grafico_forrajero(disp, ev)
            if fig_forrajero:
                try:
                    img_bytes = fig_forrajero.to_image(format='png', width=1000, height=700, scale=2)
                    img_path = os.path.join(tmpdir, 'forrajero.png')
                    with open(img_path, 'wb') as f:
                        f.write(img_bytes)
                    doc.add_picture(img_path, width=Inches(6))
                    doc.add_paragraph()
                except:
                    pass

        doc.add_heading('5.1 Interpretación técnica', level=2)
        analisis_forrajero = generar_analisis_forrajero(df, stats)
        doc.add_paragraph(analisis_forrajero)

        # 6. Mapas de calor
        if sistema_mapas:
            doc.add_heading('6. MAPAS DE CALOR CONTINUOS', level=1)
            variables = ['carbono', 'ndvi', 'ndwi', 'biodiversidad', 'forraje']
            titulos = ['Carbono (ton C/ha)', 'NDVI', 'NDWI', 'Biodiversidad (Shannon)', 'Productividad Forrajera (kg MS/ha)']
            for var, tit in zip(variables, titulos):
                mapa = sistema_mapas.crear_mapa_estatico(resultados, var, gdf)
                if mapa:
                    doc.add_heading(tit, level=2)
                    img_path = os.path.join(tmpdir, f'mapa_{var}.png')
                    with open(img_path, 'wb') as f:
                        f.write(mapa.getvalue())
                    doc.add_picture(img_path, width=Inches(6))
                    doc.add_paragraph()

        # 7. Recomendaciones Integradas
        doc.add_heading('7. RECOMENDACIONES DE MANEJO', level=1)
        recomendaciones = generar_recomendaciones_integradas(df, stats)
        doc.add_paragraph(recomendaciones)

        # 8. Metadatos
        doc.add_heading('8. METADATOS', level=1)
        metadatos = [
            ('Generado por', 'Sistema Satelital de Análisis Ambiental v3.0 con IA Groq'),
            ('Fecha de generación', datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
            ('Número de puntos', str(stats['num_puntos']))
        ]
        for key, val in metadatos:
            p = doc.add_paragraph()
            p.add_run(f"{key}: ").bold = True
            p.add_run(val)

        docx_output = BytesIO()
        doc.save(docx_output)
        docx_output.seek(0)
        return docx_output

# ===============================
# FUNCIONES AUXILIARES
# ===============================
def validar_y_corregir_crs(gdf):
    if gdf is None or len(gdf) == 0:
        return gdf
    try:
        if gdf.crs is None:
            gdf = gdf.set_crs('EPSG:4326', inplace=False)
            st.info("ℹ️ Se asignó EPSG:4326 al archivo (no tenía CRS)")
        elif str(gdf.crs).upper() != 'EPSG:4326':
            original_crs = str(gdf.crs)
            gdf = gdf.to_crs('EPSG:4326')
            st.info(f"ℹ️ Transformado de {original_crs} a EPSG:4326")
        return gdf
    except Exception as e:
        st.warning(f"⚠️ Error al corregir CRS: {str(e)}")
        return gdf

def calcular_superficie(gdf):
    try:
        if gdf is None or len(gdf) == 0:
            return 0.0
        gdf = validar_y_corregir_crs(gdf)
        bounds = gdf.total_bounds
        if bounds[0] < -180 or bounds[2] > 180 or bounds[1] < -90 or bounds[3] > 90:
            st.warning("⚠️ Coordenadas fuera de rango para cálculo preciso de área")
            area_grados2 = gdf.geometry.area.sum()
            area_m2 = area_grados2 * 111000 * 111000
            return area_m2 / 10000
        gdf_projected = gdf.to_crs('EPSG:3857')
        area_m2 = gdf_projected.geometry.area.sum()
        return area_m2 / 10000
    except Exception as e:
        try:
            return gdf.geometry.area.sum() / 10000
        except:
            return 0.0

def dividir_poligono_en_cuadricula(poligono, puntos_forraje, n_celdas=100):
    try:
        bounds = poligono.bounds
        minx, miny, maxx, maxy = bounds
        n_cols = int(np.sqrt(n_celdas * (maxx - minx) / (maxy - miny)))
        n_rows = int(n_celdas / n_cols)
        if n_rows == 0:
            n_rows = 1
        width = (maxx - minx) / n_cols
        height = (maxy - miny) / n_rows
        celdas = []
        productividades = []
        for i in range(n_rows):
            for j in range(n_cols):
                cell_minx = minx + j * width
                cell_maxx = minx + (j + 1) * width
                cell_miny = miny + i * height
                cell_maxy = miny + (i + 1) * height
                cell_poly = Polygon([(cell_minx, cell_miny), (cell_maxx, cell_miny), (cell_maxx, cell_maxy), (cell_minx, cell_maxy)])
                intersection = poligono.intersection(cell_poly)
                if intersection.is_empty or intersection.area == 0:
                    continue
                puntos_dentro = [p['productividad_kg_ms_ha'] for p in puntos_forraje if Point(p['lon'], p['lat']).within(intersection)]
                if puntos_dentro:
                    prod_promedio = np.mean(puntos_dentro)
                else:
                    min_dist = float('inf')
                    prod_cercano = None
                    for p in puntos_forraje:
                        point = Point(p['lon'], p['lat'])
                        dist = intersection.distance(point)
                        if dist < min_dist:
                            min_dist = dist
                            prod_cercano = p['productividad_kg_ms_ha']
                    prod_promedio = prod_cercano if prod_cercano is not None else 0
                celdas.append(intersection)
                productividades.append(prod_promedio)
        gdf_celdas = gpd.GeoDataFrame({'geometry': celdas, 'productividad_kg_ms_ha': productividades}, crs='EPSG:4326')
        return gdf_celdas
    except Exception as e:
        st.warning(f"Error en dividir cuadrícula: {str(e)}")
        return gpd.GeoDataFrame()

def cargar_shapefile_desde_zip(zip_file):
    try:
        with tempfile.TemporaryDirectory() as tmp_dir:
            with zipfile.ZipFile(zip_file, 'r') as zip_ref:
                zip_ref.extractall(tmp_dir)
            shp_files = [f for f in os.listdir(tmp_dir) if f.endswith('.shp')]
            if shp_files:
                shp_path = os.path.join(tmp_dir, shp_files[0])
                gdf = gpd.read_file(shp_path)
                gdf = validar_y_corregir_crs(gdf)
                return gdf
            else:
                st.error("❌ No se encontró ningún archivo .shp en el ZIP")
                return None
    except Exception as e:
        st.error(f"❌ Error cargando shapefile desde ZIP: {str(e)}")
        return None

def parsear_kml_manual(contenido_kml):
    try:
        root = ET.fromstring(contenido_kml)
        namespaces = {'kml': 'http://www.opengis.net/kml/2.2'}
        polygons = []
        for polygon_elem in root.findall('.//kml:Polygon', namespaces):
            coords_elem = polygon_elem.find('.//kml:coordinates', namespaces)
            if coords_elem is not None and coords_elem.text:
                coord_text = coords_elem.text.strip()
                coord_list = []
                for coord_pair in coord_text.split():
                    parts = coord_pair.split(',')
                    if len(parts) >= 2:
                        lon = float(parts[0])
                        lat = float(parts[1])
                        coord_list.append((lon, lat))
                if len(coord_list) >= 3:
                    polygons.append(Polygon(coord_list))
        if not polygons:
            for multi_geom in root.findall('.//kml:MultiGeometry', namespaces):
                for polygon_elem in multi_geom.findall('.//kml:Polygon', namespaces):
                    coords_elem = polygon_elem.find('.//kml:coordinates', namespaces)
                    if coords_elem is not None and coords_elem.text:
                        coord_text = coords_elem.text.strip()
                        coord_list = []
                        for coord_pair in coord_text.split():
                            parts = coord_pair.split(',')
                            if len(parts) >= 2:
                                lon = float(parts[0])
                                lat = float(parts[1])
                                coord_list.append((lon, lat))
                        if len(coord_list) >= 3:
                            polygons.append(Polygon(coord_list))
        if polygons:
            gdf = gpd.GeoDataFrame({'geometry': polygons}, crs='EPSG:4326')
            return gdf
        else:
            for placemark in root.findall('.//kml:Placemark', namespaces):
                for elem_name in ['Polygon', 'LineString', 'Point', 'LinearRing']:
                    elem = placemark.find(f'.//kml:{elem_name}', namespaces)
                    if elem is not None:
                        coords_elem = elem.find('.//kml:coordinates', namespaces)
                        if coords_elem is not None and coords_elem.text:
                            coord_text = coords_elem.text.strip()
                            coord_list = []
                            for coord_pair in coord_text.split():
                                parts = coord_pair.split(',')
                                if len(parts) >= 2:
                                    lon = float(parts[0])
                                    lat = float(parts[1])
                                    coord_list.append((lon, lat))
                            if len(coord_list) >= 3:
                                polygons.append(Polygon(coord_list))
                            break
        if polygons:
            gdf = gpd.GeoDataFrame({'geometry': polygons}, crs='EPSG:4326')
            return gdf
        return None
    except Exception as e:
        st.error(f"❌ Error parseando KML manualmente: {str(e)}")
        return None

def cargar_kml(kml_file):
    try:
        if kml_file.name.endswith('.kmz'):
            with tempfile.TemporaryDirectory() as tmp_dir:
                with zipfile.ZipFile(kml_file, 'r') as zip_ref:
                    zip_ref.extractall(tmp_dir)
                kml_files = [f for f in os.listdir(tmp_dir) if f.endswith('.kml')]
                if kml_files:
                    kml_path = os.path.join(tmp_dir, kml_files[0])
                    with open(kml_path, 'r', encoding='utf-8') as f:
                        contenido = f.read()
                    gdf = parsear_kml_manual(contenido)
                    if gdf is not None:
                        return gdf
                    else:
                        try:
                            gdf = gpd.read_file(kml_path)
                            gdf = validar_y_corregir_crs(gdf)
                            return gdf
                        except:
                            st.error("❌ No se pudo cargar el archivo KML/KMZ")
                            return None
                else:
                    st.error("❌ No se encontró ningún archivo .kml en el KMZ")
                    return None
        else:
            contenido = kml_file.read().decode('utf-8')
            gdf = parsear_kml_manual(contenido)
            if gdf is not None:
                return gdf
            else:
                kml_file.seek(0)
                gdf = gpd.read_file(kml_file)
                gdf = validar_y_corregir_crs(gdf)
                return gdf
    except Exception as e:
        st.error(f"❌ Error cargando archivo KML/KMZ: {str(e)}")
        return None

def cargar_archivo_parcela(uploaded_file):
    try:
        if uploaded_file.name.endswith('.zip'):
            gdf = cargar_shapefile_desde_zip(uploaded_file)
        elif uploaded_file.name.endswith(('.kml', '.kmz')):
            gdf = cargar_kml(uploaded_file)
        elif uploaded_file.name.endswith('.geojson'):
            gdf = gpd.read_file(uploaded_file)
            gdf = validar_y_corregir_crs(gdf)
        else:
            st.error("❌ Formato de archivo no soportado")
            return None

        if gdf is not None:
            gdf = validar_y_corregir_crs(gdf)
            gdf = gdf.explode(ignore_index=True)
            gdf = gdf[gdf.geometry.geom_type.isin(['Polygon', 'MultiPolygon'])]
            if len(gdf) == 0:
                st.error("❌ No se encontraron polígonos en el archivo")
                return None
            geometria_unida = gdf.unary_union
            gdf_unido = gpd.GeoDataFrame([{'geometry': geometria_unida}], crs='EPSG:4326')
            gdf_unido = validar_y_corregir_crs(gdf_unido)
            st.info(f"✅ Se unieron {len(gdf)} polígono(s) en una sola geometría.")
            gdf_unido['id_zona'] = 1
            return gdf_unido
        return gdf
    except Exception as e:
        st.error(f"❌ Error cargando archivo: {str(e)}")
        import traceback
        st.error(f"Detalle: {traceback.format_exc()}")
        return None

# ===============================
# FUNCIÓN PRINCIPAL DE ANÁLISIS
# ===============================
def ejecutar_analisis_completo(gdf, tipo_ecosistema, num_puntos, usar_gee=False):
    try:
        area_total = calcular_superficie(gdf)
        poligono = gdf.geometry.iloc[0]
        bounds = poligono.bounds

        clima = ConectorClimaticoTropical()
        verra = MetodologiaVerra()
        biodiversidad = AnalisisBiodiversidad()
        forrajero = AnalisisForrajero()

        # Asignar sistema forrajero según ecosistema
        if tipo_ecosistema in ['pampa', 'seco', 'espinal', 'patagonico']:
            sistema_forrajero = 'pastizal_natural'
        elif tipo_ecosistema in ['amazonia', 'choco', 'yungas', 'paranaense']:
            sistema_forrajero = 'silvopastoril'
        elif tipo_ecosistema in ['monte']:
            sistema_forrajero = 'monte'
        else:
            sistema_forrajero = 'pastizal_natural'

        puntos_carbono = []
        puntos_biodiversidad = []
        puntos_ndvi = []
        puntos_ndwi = []
        puntos_ndre = []
        puntos_msavi = []
        puntos_evi = []
        puntos_forraje = []

        carbono_total = 0
        co2_total = 0
        shannon_promedio = 0
        ndvi_promedio = 0
        ndwi_promedio = 0
        area_por_punto = max(area_total / num_puntos, 0.1)

        puntos_generados = 0
        max_intentos = num_puntos * 10

        while puntos_generados < num_puntos and len(puntos_carbono) < max_intentos:
            lat = bounds[1] + random.random() * (bounds[3] - bounds[1])
            lon = bounds[0] + random.random() * (bounds[2] - bounds[0])
            point = Point(lon, lat)

            if poligono.contains(point):
                datos_clima = clima.obtener_datos_climaticos(lat, lon)
                ndvi = 0.5 + random.uniform(-0.2, 0.3)
                base_ndwi = 0.1
                if datos_clima['precipitacion'] > 2000:
                    base_ndwi += 0.3
                elif datos_clima['precipitacion'] < 800:
                    base_ndwi -= 0.2
                ndwi = base_ndwi + random.uniform(-0.2, 0.2)
                ndwi = max(-0.5, min(0.8, ndwi))
                ndre = min(1.0, max(-1.0, ndvi * 0.95 + random.uniform(-0.05, 0.1)))
                msavi = min(1.0, max(0.0, ndvi * 0.85 + random.uniform(-0.1, 0.05)))
                evi = min(1.0, max(0.0, ndvi * 1.2 + random.uniform(-0.1, 0.1)))

                carbono_info = verra.calcular_carbono_hectarea(ndvi, tipo_ecosistema, datos_clima['precipitacion'])
                biodiv_info = biodiversidad.calcular_shannon(ndvi, tipo_ecosistema, area_por_punto, datos_clima['precipitacion'])
                forraje_info = forrajero.estimar_disponibilidad_forrajera(ndvi, sistema_forrajero, area_por_punto)

                carbono_total += carbono_info['carbono_total_ton_ha'] * area_por_punto
                co2_total += carbono_info['co2_equivalente_ton_ha'] * area_por_punto
                shannon_promedio += biodiv_info['indice_shannon']
                ndvi_promedio += ndvi
                ndwi_promedio += ndwi

                puntos_carbono.append({'lat': lat, 'lon': lon, 'carbono_ton_ha': carbono_info['carbono_total_ton_ha'], 'ndvi': ndvi, 'precipitacion': datos_clima['precipitacion']})
                biodiv_info['lat'] = lat
                biodiv_info['lon'] = lon
                puntos_biodiversidad.append(biodiv_info)
                puntos_ndvi.append({'lat': lat, 'lon': lon, 'ndvi': ndvi})
                puntos_ndwi.append({'lat': lat, 'lon': lon, 'ndwi': ndwi})
                puntos_ndre.append({'lat': lat, 'lon': lon, 'ndre': ndre})
                puntos_msavi.append({'lat': lat, 'lon': lon, 'msavi': msavi})
                puntos_evi.append({'lat': lat, 'lon': lon, 'evi': evi})
                puntos_forraje.append({'lat': lat, 'lon': lon, 'productividad_kg_ms_ha': forraje_info['productividad_kg_ms_ha']})

                puntos_generados += 1

        if puntos_generados > 0:
            shannon_promedio /= puntos_generados
            ndvi_promedio /= puntos_generados
            ndwi_promedio /= puntos_generados

        carbono_promedio = verra.calcular_carbono_hectarea(ndvi_promedio, tipo_ecosistema, 1500)

        # Análisis forrajero
        disponibilidad_forrajera = forrajero.estimar_disponibilidad_forrajera(ndvi_promedio, sistema_forrajero, area_total)
        equivalentes_vaca = forrajero.calcular_equivalentes_vaca(disponibilidad_forrajera['forraje_aprovechable_kg_ms'], dias_permanencia=30)
        sublotes = forrajero.dividir_lote_en_sublotes(area_total, disponibilidad_forrajera['productividad_kg_ms_ha'], heterogeneidad=0.3)
        gdf_cuadricula = dividir_poligono_en_cuadricula(poligono, puntos_forraje, n_celdas=200)

        # === PASTOREO RACIONAL VOISIN (PRV) ===
        prv = ModeloPRV()
        prv_temperatura = 22
        prv_precipitacion = 100
        prv_descanso = prv.calcular_periodo_descanso(prv_temperatura, prv_precipitacion)
        prv_ocupacion = prv.calcular_periodo_ocupacion(
            disponibilidad_forrajera['productividad_kg_ms_ha'],
            num_ev=equivalentes_vaca['ev_recomendado'],
            area_ha=area_total / max(len(sublotes), 1)
        )
        prv_num_potreros = prv.calcular_numero_potreros(prv_descanso, prv_ocupacion)

        # Generar potreros desde cuadrícula forrajera
        if gdf_cuadricula is not None and not gdf_cuadricula.empty:
            gdf_potreros = prv.generar_potreros_desde_cuadricula(
                gdf_cuadricula, num_potreros=prv_num_potreros
            )
            plan_pastoreo = prv.planificar_ciclo_prv(
                num_potreros=len(gdf_potreros) if gdf_potreros is not None else prv_num_potreros,
                descanso_dias=prv_descanso,
                ocupacion_dias=prv_ocupacion,
                num_ciclos=4
            )
            estados_prv_df = prv.calcular_estado_potreros(
                gdf_potreros, plan_pastoreo
            ) if gdf_potreros is not None else pd.DataFrame()
            resumen_prv = prv.resumen_ejecutivo(
                gdf_potreros, estados_prv_df, plan_pastoreo
            ) if gdf_potreros is not None else {}
            recomendaciones_prv = prv.generar_recomendaciones_prv(
                gdf_potreros, estados_prv_df, prv_temperatura, prv_precipitacion
            ) if gdf_potreros is not None else "No hay datos suficientes para PRV."
        else:
            gdf_potreros = None
            plan_pastoreo = pd.DataFrame()
            estados_prv_df = pd.DataFrame()
            resumen_prv = {}
            recomendaciones_prv = "No hay datos suficientes para PRV."

        resultados = {
            'area_total_ha': area_total,
            'carbono_total_ton': round(carbono_total, 2),
            'co2_total_ton': round(co2_total, 2),
            'carbono_promedio_ha': round(carbono_total / area_total, 2) if area_total > 0 else 0,
            'shannon_promedio': round(shannon_promedio, 3),
            'ndvi_promedio': round(ndvi_promedio, 3),
            'ndwi_promedio': round(ndwi_promedio, 3),
            'puntos_carbono': puntos_carbono,
            'puntos_biodiversidad': puntos_biodiversidad,
            'puntos_ndvi': puntos_ndvi,
            'puntos_ndwi': puntos_ndwi,
            'puntos_ndre': puntos_ndre,
            'puntos_msavi': puntos_msavi,
            'puntos_evi': puntos_evi,
            'puntos_forraje': puntos_forraje,
            'gdf_cuadricula': gdf_cuadricula,
            'tipo_ecosistema': tipo_ecosistema,
            'num_puntos': puntos_generados,
            'desglose_promedio': carbono_promedio['desglose'] if carbono_promedio else {},
            'usar_gee': usar_gee,
            'analisis_forrajero': {
                'sistema_forrajero': sistema_forrajero,
                'disponibilidad_forrajera': disponibilidad_forrajera,
                'equivalentes_vaca': equivalentes_vaca,
                'sublotes': sublotes,
                'forrajero': forrajero
            },
            'prv': {
                'modelo': prv,
                'gdf_potreros': gdf_potreros,
                'plan_pastoreo': plan_pastoreo,
                'estados_prv': estados_prv_df,
                'resumen': resumen_prv,
                'recomendaciones': recomendaciones_prv,
                'descanso_dias': prv_descanso,
                'ocupacion_dias': prv_ocupacion,
                'num_potreros': prv_num_potreros,
                'temperatura': prv_temperatura,
                'precipitacion': prv_precipitacion,
            }
        }
        return resultados
    except Exception as e:
        st.error(f"Error en ejecutar_analisis_completo: {str(e)}")
        import traceback
        st.error(traceback.format_exc())
        return None

# ===============================
# FUNCIONES DE VISUALIZACIÓN
# ===============================
def mostrar_mapas_calor():
    st.header("🗺️ Mapas de Calor Continuos")
    if st.session_state.poligono_data is None or st.session_state.resultados is None:
        st.info("Ejecute el análisis primero.")
        return

    tab1, tab2, tab3, tab4, tab5, tab6, tab7 = st.tabs([
        "🌍 Área Base", "🌳 Carbono", "📈 NDVI", "💧 NDWI", "🦋 Biodiversidad", "🌿 Forrajero", "🎭 Combinado"
    ])

    with tab1:
        st.subheader("Mapa Base del Área de Estudio")
        if st.session_state.mapa:
            folium_static(st.session_state.mapa, width=1000, height=650)
        else:
            st.info("No hay mapa base.")

    sistema = SistemaMapas()
    with tab2:
        if 'puntos_carbono' in st.session_state.resultados:
            mapa = sistema.crear_mapa_calor_interpolado(st.session_state.resultados, 'carbono', st.session_state.poligono_data)
            if mapa:
                folium_static(mapa, width=1000, height=650)
            else:
                st.warning("No se pudo generar el mapa.")
    with tab3:
        if 'puntos_ndvi' in st.session_state.resultados:
            mapa = sistema.crear_mapa_calor_interpolado(st.session_state.resultados, 'ndvi', st.session_state.poligono_data)
            if mapa:
                folium_static(mapa, width=1000, height=650)
    with tab4:
        if 'puntos_ndwi' in st.session_state.resultados:
            mapa = sistema.crear_mapa_calor_interpolado(st.session_state.resultados, 'ndwi', st.session_state.poligono_data)
            if mapa:
                folium_static(mapa, width=1000, height=650)
    with tab5:
        if 'puntos_biodiversidad' in st.session_state.resultados:
            mapa = sistema.crear_mapa_calor_interpolado(st.session_state.resultados, 'biodiversidad', st.session_state.poligono_data)
            if mapa:
                folium_static(mapa, width=1000, height=650)
    with tab6:
        if 'puntos_forraje' in st.session_state.resultados:
            mapa = sistema.crear_mapa_calor_interpolado(st.session_state.resultados, 'forraje', st.session_state.poligono_data)
            if mapa:
                folium_static(mapa, width=1000, height=650)
    with tab7:
        st.subheader("🎭 Mapa Combinado - Todas las Capas")
        if st.session_state.resultados:
            mapa_combinado = sistema.crear_mapa_combinado_interpolado(st.session_state.resultados, st.session_state.poligono_data)
            if mapa_combinado:
                folium_static(mapa_combinado, width=1000, height=650)
                st.info("Use el control de capas en la esquina superior derecha para activar/desactivar cada variable.")
            else:
                st.warning("No se pudo generar el mapa combinado.")
        else:
            st.info("Ejecute el análisis primero para ver el mapa combinado")

def mostrar_dashboard():
    st.header("📊 Dashboard Ejecutivo")
    if st.session_state.resultados is None:
        st.info("Ejecute el análisis primero.")
        return
    res = st.session_state.resultados

    # Extraer métricas forrajeras si existen
    forraje_kg = 0
    ev_recomendado = 0
    if 'analisis_forrajero' in res:
        forraje_kg = res['analisis_forrajero']['disponibilidad_forrajera']['productividad_kg_ms_ha']
        ev_recomendado = res['analisis_forrajero']['equivalentes_vaca']['ev_recomendado']

    html_kpi = Visualizaciones.crear_metricas_kpi(
        res.get('carbono_total_ton', 0),
        res.get('co2_total_ton', 0),
        res.get('shannon_promedio', 0),
        res.get('area_total_ha', 0),
        ndvi=res.get('ndvi_promedio', 0),
        ndwi=res.get('ndwi_promedio', 0),
        forraje_kg=forraje_kg,
        ev=ev_recomendado,
    )
    st.markdown(html_kpi, unsafe_allow_html=True)

    eco = str(res.get('tipo_ecosistema', '')).title()
    pts = str(res.get('num_puntos', 0))
    sis = str(res.get('analisis_forrajero', {}).get('sistema_forrajero', 'N/A')).replace('_', ' ').title()
    st.markdown(
        '<div style="display:flex;gap:1rem;align-items:center;margin:1rem 0 1.5rem;'
        'background:rgba(255,255,255,0.03);border-radius:12px;padding:0.75rem 1.25rem;'
        'border:1px solid rgba(255,255,255,0.05);">'
        '<span style="font-size:0.85rem;color:#94a3b8;">Ecosistema:</span>'
        '<span style="font-weight:600;color:#e2e8f0;">' + eco + '</span>'
        '<span style="color:#475569;">|</span>'
        '<span style="font-size:0.85rem;color:#94a3b8;">Puntos de muestreo:</span>'
        '<span style="font-weight:600;color:#e2e8f0;">' + pts + '</span>'
        '<span style="color:#475569;">|</span>'
        '<span style="font-size:0.85rem;color:#94a3b8;">Sistema forrajero:</span>'
        '<span style="font-weight:600;color:#e2e8f0;">' + sis + '</span></div>',
        unsafe_allow_html=True
    )

    col_left, col_right = st.columns([1.2, 1])

    with col_left:
        st.markdown('<h3 style="display:flex;align-items:center;gap:0.5rem;margin-bottom:0.5rem;">🌳 Distribución de Carbono</h3>', unsafe_allow_html=True)
        fig_carbono = Visualizaciones.crear_grafico_barras_carbono(res.get('desglose_promedio', {}))
        if fig_carbono:
            fig_carbono.update_layout(
                paper_bgcolor='rgba(0,0,0,0)', plot_bgcolor='rgba(0,0,0,0)',
                font=dict(color='#cbd5e1', size=11),
                title=None,
            )
            fig_carbono.update_xaxes(gridcolor='rgba(255,255,255,0.05)')
            fig_carbono.update_yaxes(gridcolor='rgba(255,255,255,0.05)')
            st.plotly_chart(fig_carbono, use_container_width=True)

    with col_right:
        st.markdown('<h3 style="display:flex;align-items:center;gap:0.5rem;margin-bottom:0.5rem;">🦋 Perfil de Biodiversidad</h3>', unsafe_allow_html=True)
        if res.get('puntos_biodiversidad'):
            fig_biodiv = Visualizaciones.crear_grafico_radar_biodiversidad(res['puntos_biodiversidad'][0])
            if fig_biodiv:
                fig_biodiv.update_layout(
                    paper_bgcolor='rgba(0,0,0,0)', plot_bgcolor='rgba(0,0,0,0)',
                    font=dict(color='#cbd5e1', size=11),
                title=None,
                    polar=dict(bgcolor='rgba(0,0,0,0)',
                               radialaxis=dict(gridcolor='rgba(255,255,255,0.1)', linecolor='rgba(255,255,255,0.05)')),
                )
                st.plotly_chart(fig_biodiv, use_container_width=True)

    # Tabla resumen
    if 'analisis_forrajero' in res:
        prod = f"{forraje_kg:,.0f}"
        apro = f"{res['analisis_forrajero']['disponibilidad_forrajera']['forraje_aprovechable_kg_ms']/1000:,.1f}"
        evs = f"{ev_recomendado:.1f}"
        subs = str(len(res['analisis_forrajero']['sublotes']))
        st.markdown(
            '<div style="display:grid;grid-template-columns:repeat(4,1fr);gap:0.75rem;margin:1.5rem 0 0.5rem;'
            'padding:1rem;background:rgba(255,255,255,0.02);border-radius:12px;'
            'border:1px solid rgba(255,255,255,0.04);">'
            '<div><span style="color:#64748b;font-size:0.7rem;text-transform:uppercase;letter-spacing:0.05em;">Productividad</span>'
            '<br><span style="color:#e2e8f0;font-size:1.1rem;font-weight:600;">' + prod + '</span>'
            '<span style="color:#475569;font-size:0.75rem;"> kg MS/ha</span></div>'
            '<div><span style="color:#64748b;font-size:0.7rem;text-transform:uppercase;letter-spacing:0.05em;">Forraje Aprovechable</span>'
            '<br><span style="color:#e2e8f0;font-size:1.1rem;font-weight:600;">' + apro + '</span>'
            '<span style="color:#475569;font-size:0.75rem;"> ton MS</span></div>'
            '<div><span style="color:#64748b;font-size:0.7rem;text-transform:uppercase;letter-spacing:0.05em;">EV Recomendados</span>'
            '<br><span style="color:#e2e8f0;font-size:1.1rem;font-weight:600;">' + evs + '</span>'
            '<span style="color:#475569;font-size:0.75rem;"> EV</span></div>'
            '<div><span style="color:#64748b;font-size:0.7rem;text-transform:uppercase;letter-spacing:0.05em;">Sublotes / Potreros</span>'
            '<br><span style="color:#e2e8f0;font-size:1.1rem;font-weight:600;">' + subs + '</span>'
            '<span style="color:#475569;font-size:0.75rem;"> sublotes</span></div></div>',
            unsafe_allow_html=True
        )

def mostrar_carbono():
    st.header("🌳 Análisis de Carbono")
    if st.session_state.resultados is None:
        st.info("Ejecute el análisis primero.")
        return
    res = st.session_state.resultados
    co2_ton = res.get('co2_total_ton', 0)
    c_ton = res.get('carbono_total_ton', 0)

    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("Carbono Total", f"{c_ton:,.0f} ton C")
    with col2:
        st.metric("CO₂ Secuestrado", f"{co2_ton:,.0f} ton CO₂e")
    with col3:
        precio_ton = 15
        valor_economico = co2_ton * precio_ton
        st.metric("Valor Aprox.", f"${valor_economico:,.0f} USD")

    if res.get('desglose_promedio'):
        st.subheader("📊 Distribución por Pools de Carbono")
        df_pools = pd.DataFrame({
            'Pool': list(res['desglose_promedio'].keys()),
            'Ton C/ha': list(res['desglose_promedio'].values())
        })
        fig_pools = go.Figure(data=[go.Bar(
            x=df_pools['Pool'], y=df_pools['Ton C/ha'],
            marker=dict(color=['#065f46', '#059669', '#10b981', '#34d399', '#6ee7b7']),
            text=df_pools['Ton C/ha'].apply(lambda x: f'{x:.1f}'),
            textposition='outside',
        )])
        fig_pools.update_layout(
            title='', xaxis_title='Pool', yaxis_title='Ton C/ha',
            height=350, paper_bgcolor='rgba(0,0,0,0)', plot_bgcolor='rgba(0,0,0,0)',
            font=dict(color='#cbd5e1', size=11),
            xaxis=dict(gridcolor='rgba(255,255,255,0.05)'),
            yaxis=dict(gridcolor='rgba(255,255,255,0.05)'),
        )
        st.plotly_chart(fig_pools, use_container_width=True)

    # =================================================================
    # CALCULADORA DE HUELLA DE CARBONO GANADERA
    # =================================================================
    st.divider()
    st.markdown(
        '<div style="display:flex;align-items:center;gap:0.6rem;margin-bottom:0.5rem;">'
        '<span style="font-size:1.5rem;">🐄</span>'
        '<span style="font-size:1.15rem;font-weight:600;color:#f1f5f9;">Calculadora de Huella de Carbono Ganadera</span></div>'
        '<p style="color:#94a3b8;font-size:0.85rem;margin-bottom:1.25rem;">'
        'Estimá las emisiones de tu rodeo y comparalas con el carbono secuestrado en tu campo'
        ' para conocer tu <strong>balance de carbono</strong> y el potencial de créditos comercializables.</p>',
        unsafe_allow_html=True
    )

    with st.expander("📝 Datos del rodeo", expanded=True):
        col1, col2 = st.columns(2)
        with col1:
            tipo_ganado = st.selectbox("Tipo de ganado", ["bovino carne", "bovino leche", "ovina", "mixto"], index=0)
            cabezas = st.number_input("Cantidad de cabezas", min_value=1, max_value=100000, value=100, step=10)
            peso_prom = st.number_input("Peso vivo promedio (kg)", min_value=50, max_value=1200, value=400, step=10)
        with col2:
            sistema_alimentacion = st.selectbox("Sistema de alimentación", [
                "pastoreo permanente", "pastoreo + suplemento", "feedlot / confinamiento", "mixto"
            ], index=0)
            categoria = st.selectbox("Categoría predominante", [
                "vacas adultas", "novillos", "terneros/as", "cria (vaca + ternero)", "ciclo completo"
            ], index=0)
            manejo_estiércol = st.selectbox("Manejo de estiércol", [
                "pastoreo (deposición en campo)", "recolección + compostaje",
                "recolección + laguna", "digestor anaeróbico"
            ], index=0)

        co2_secuestrado = co2_ton  # ton CO2e del análisis satelital

        usar_ia = st.checkbox("Usar IA (Groq) para recomendaciones detalladas", value=False,
                              help="Si no hay API key configurada, se usará una estimación automática.")

        if st.button("🔥 Calcular huella y balance", use_container_width=True, type="primary"):
            with st.spinner("Calculando emisiones y balance de carbono..."):
                # Factores de emisión IPCC (Tier 1)
                # CH4 entérico: kg CH4/cabeza/año
                if tipo_ganado == "bovino carne":
                    if categoria == "vacas adultas":
                        factor_enterico = 62  # kg CH4/año
                    elif categoria == "novillos":
                        factor_enterico = 47
                    elif categoria == "terneros/as":
                        factor_enterico = 25
                    else:  # cria o ciclo completo
                        factor_enterico = 55
                elif tipo_ganado == "bovino leche":
                    factor_enterico = 85
                elif tipo_ganado == "ovina":
                    factor_enterico = 8
                else:
                    factor_enterico = 50

                # Ajuste por peso
                factor_enterico = factor_enterico * (peso_prom / 400)

                # CH4 manure: kg CH4/cabeza/año
                if manejo_estiércol == "pastoreo (deposición en campo)":
                    factor_manure_ch4 = 1.5
                    factor_manure_n2o = 0.02
                elif manejo_estiércol == "recolección + compostaje":
                    factor_manure_ch4 = 3.0
                    factor_manure_n2o = 0.005
                elif manejo_estiércol == "recolección + laguna":
                    factor_manure_ch4 = 18.0
                    factor_manure_n2o = 0.01
                else:  # digestor
                    factor_manure_ch4 = 2.0
                    factor_manure_n2o = 0.005

                # Ajuste por alimentación
                if sistema_alimentacion == "pastoreo permanente":
                    factor_enterico *= 1.0
                    factor_manure_ch4 *= 0.8
                elif sistema_alimentacion == "pastoreo + suplemento":
                    factor_enterico *= 0.95
                    factor_manure_ch4 *= 0.9
                elif sistema_alimentacion == "feedlot / confinamiento":
                    factor_enterico *= 0.85
                    factor_manure_ch4 *= 1.8
                else:  # mixto
                    factor_enterico *= 0.92
                    factor_manure_ch4 *= 1.0

                # Cálculos
                ch4_enterico_kg = factor_enterico * cabezas
                ch4_manure_kg = factor_manure_ch4 * cabezas
                ch4_total_kg = ch4_enterico_kg + ch4_manure_kg
                ch4_total_co2e = ch4_total_kg * 28  # GWP CH4 = 28

                n2o_kg = factor_manure_n2o * cabezas * (peso_prom / 400)
                n2o_total_co2e = n2o_kg * 265  # GWP N2O = 265

                emisiones_totales_co2e = ch4_total_co2e + n2o_total_co2e
                balance_co2e = co2_secuestrado - emisiones_totales_co2e

            # === RESULTADOS ===
            st.subheader("📊 Resultados del Balance de Carbono")

            col1, col2, col3, col4 = st.columns(4)
            with col1:
                st.metric("CH₄ entérico", f"{ch4_enterico_kg:,.0f} kg CH₄/año")
            with col2:
                st.metric("CH₄ estiércol", f"{ch4_manure_kg:,.0f} kg CH₄/año")
            with col3:
                st.metric("N₂O estiércol", f"{n2o_kg:.1f} kg N₂O/año")
            with col4:
                st.metric("Emisiones totales", f"{emisiones_totales_co2e:,.0f} ton CO₂e/año",
                          delta=f"{emisiones_totales_co2e / max(co2_secuestrado, 1) * 100:.0f}% del secuestro")

            # Delta
            if balance_co2e > 0:
                delta_color = "green"
                delta_texto = "positivo — tu campo captura más carbono del que emite 🎉"
            else:
                delta_color = "red"
                delta_texto = "negativo — las emisiones superan la captura. Revisá las recomendaciones abajo."

            color_balance = '#10b981' if balance_co2e > 0 else '#ef4444'
            color_balance_text = '#6ee7b7' if balance_co2e > 0 else '#fca5a5'
            st.markdown(
                '<div style="background:rgba(255,255,255,0.03);border:1px solid rgba(255,255,255,0.06);'
                'border-radius:14px;padding:1.5rem;margin:1rem 0;">'
                '<div style="display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:1rem;">'
                '<div><div style="color:#94a3b8;font-size:0.75rem;text-transform:uppercase;letter-spacing:0.05em;">CO\u2082 Secuestrado (sat\u00e9lite)</div>'
                '<div style="font-size:2rem;font-weight:700;color:#10b981;">' + f'{co2_secuestrado:,.0f}' + ' <span style="font-size:0.9rem;color:#6ee7b7;">ton CO\u2082e/a\u00f1o</span></div></div>'
                '<div style="color:#475569;font-size:1.5rem;">\u2212</div>'
                '<div><div style="color:#94a3b8;font-size:0.75rem;text-transform:uppercase;letter-spacing:0.05em;">Emisiones ganaderas</div>'
                '<div style="font-size:2rem;font-weight:700;color:#ef4444;">' + f'{emisiones_totales_co2e:,.0f}' + ' <span style="font-size:0.9rem;color:#fca5a5;">ton CO\u2082e/a\u00f1o</span></div></div>'
                '<div style="color:#475569;font-size:1.5rem;">=</div>'
                '<div><div style="color:#94a3b8;font-size:0.75rem;text-transform:uppercase;letter-spacing:0.05em;">Balance neto</div>'
                '<div style="font-size:2rem;font-weight:700;color:' + color_balance + ';">' + f'{balance_co2e:+,.0f}' + ' <span style="font-size:0.9rem;color:' + color_balance_text + ';">ton CO\u2082e/a\u00f1o</span></div></div></div>'
                '<div style="margin-top:1rem;padding:0.5rem 0.75rem;background:rgba(255,255,255,0.04);border-radius:8px;font-size:0.85rem;color:#cbd5e1;">'
                '\U0001f4cc Balance ' + delta_texto + '</div></div>',
                unsafe_allow_html=True
            )

            # Potencial económico
            st.subheader("💰 Potencial de Créditos de Carbono")
            if balance_co2e > 0:
                creditos_vendibles = balance_co2e * 0.7  # 70% comercializable (margen de seguridad)
                ingreso_anual = creditos_vendibles * precio_ton
                ingreso_10a = ingreso_anual * 10

                col1, col2, col3 = st.columns(3)
                with col1:
                    st.metric("Créditos comercializables", f"{creditos_vendibles:,.0f} ton CO₂e/año",
                              help="Estimación conservadora: 70% del balance neto")
                with col2:
                    st.metric("Ingreso anual estimado", f"${ingreso_anual:,.0f} USD/año",
                              help=f"Precio de referencia: ${precio_ton}/ton CO₂e (mercado voluntario)")
                with col3:
                    st.metric("Proyección 10 años", f"${ingreso_10a:,.0f} USD",
                              help="Estimación lineal sin ajuste por inflación ni crecimiento")
            else:
                st.warning("""
                **Balance negativo.** No hay excedente de carbono para comercializar.
                Revisá las recomendaciones abajo para reducir emisiones y aumentar la captura de carbono.
                """)

            # Gráfico de balance
            fig_balance = go.Figure()
            fig_balance.add_trace(go.Bar(
                x=['CO₂ Secuestrado', 'Emisiones Ganaderas', 'Balance Neto'],
                y=[co2_secuestrado, emisiones_totales_co2e, balance_co2e],
                marker=dict(color=['#10b981', '#ef4444', '#3b82f6']),
                text=[f'{co2_secuestrado:,.0f}', f'{emisiones_totales_co2e:,.0f}', f'{balance_co2e:+,.0f}'],
                textposition='outside',
            ))
            fig_balance.update_layout(
                title='Balance de Carbono (ton CO₂e/año)',
                height=400, paper_bgcolor='rgba(0,0,0,0)', plot_bgcolor='rgba(0,0,0,0)',
                font=dict(color='#cbd5e1', size=11),
                xaxis=dict(gridcolor='rgba(255,255,255,0.05)'),
                yaxis=dict(gridcolor='rgba(255,255,255,0.05)', title='ton CO₂e/año'),
            )
            st.plotly_chart(fig_balance, use_container_width=True)

def mostrar_biodiversidad():
    st.header("🦋 Análisis de Biodiversidad")
    if st.session_state.resultados is None:
        st.info("Ejecute el análisis primero.")
        return
    res = st.session_state.resultados
    if res.get('puntos_biodiversidad'):
        biodiv = res['puntos_biodiversidad'][0]
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Índice Shannon", f"{biodiv.get('indice_shannon', 0):.3f}", biodiv.get('categoria', ''))
        with col2:
            st.metric("Riqueza especies", f"{biodiv.get('riqueza_especies', 0)}")
        with col3:
            st.metric("Abundancia total", f"{biodiv.get('abundancia_total', 0):,}")
        # Gráfico de distribución
        shannon_vals = [p.get('indice_shannon', 0) for p in res['puntos_biodiversidad']]
        fig = go.Figure(data=[go.Histogram(x=shannon_vals, nbinsx=15, marker_color='#8b5cf6')])
        fig.update_layout(title='Distribución del Índice de Shannon', xaxis_title='Valor', yaxis_title='Frecuencia', height=400)
        st.plotly_chart(fig, use_container_width=True)

def mostrar_analisis_forrajero():
    st.header("🐮 Análisis Forrajero")
    if st.session_state.resultados is None or 'analisis_forrajero' not in st.session_state.resultados:
        st.info("Ejecute el análisis completo primero.")
        return
    res = st.session_state.resultados
    forrajero_data = res['analisis_forrajero']
    disp = forrajero_data['disponibilidad_forrajera']
    ev = forrajero_data['equivalentes_vaca']

    st.subheader("🌿 Disponibilidad Forrajera")
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("Productividad", f"{disp['productividad_kg_ms_ha']:,.0f}", "kg MS/ha")
    with col2:
        st.metric("Disponible total", f"{disp['disponibilidad_total_kg_ms']/1000:,.1f}", "ton MS")
    with col3:
        st.metric("Aprovechable", f"{disp['forraje_aprovechable_kg_ms']/1000:,.1f}", "ton MS")
    with col4:
        st.metric("Categoría", disp['categoria_productividad'].title())

    st.subheader("🐄 Equivalentes Vaca")
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("EV por día", f"{ev['ev_por_dia']:.1f}")
    with col2:
        st.metric("EV para 30 días", f"{ev['ev_para_periodo']:.1f}")
    with col3:
        st.metric("EV recomendado", f"{ev['ev_recomendado']:.1f}")

    if 'sublotes' in forrajero_data and forrajero_data['sublotes']:
        st.subheader("📋 Sublotes")
        df_sub = pd.DataFrame(forrajero_data['sublotes'])
        st.dataframe(df_sub, use_container_width=True, hide_index=True)

    # Gráfico forrajero
    fig_forrajero = Visualizaciones.crear_grafico_forrajero(disp, ev)
    if fig_forrajero:
        st.plotly_chart(fig_forrajero, use_container_width=True)

    # Mapa de sublotes (coroplético)
    if 'gdf_cuadricula' in res and not res['gdf_cuadricula'].empty:
        st.subheader("🗺️ Mapa de Productividad por Sublotes")
        sistema = SistemaMapas()
        try:
            m = SistemaMapas.crear_mapa_con_base(res['gdf_cuadricula'])
            min_prod = res['gdf_cuadricula']['productividad_kg_ms_ha'].min()
            max_prod = res['gdf_cuadricula']['productividad_kg_ms_ha'].max()
            colormap = LinearColormap(colors=['#8B4513', '#CD853F', '#F4A460', '#9ACD32', '#32CD32', '#006400'], vmin=min_prod, vmax=max_prod)
            colormap.caption = 'Productividad Forrajera (kg MS/ha)'
            folium.GeoJson(
                res['gdf_cuadricula'],
                style_function=lambda feature: {
                    'fillColor': colormap(feature['properties']['productividad_kg_ms_ha']),
                    'color': 'black',
                    'weight': 0.5,
                    'fillOpacity': 0.7
                },
                tooltip=folium.GeoJsonTooltip(fields=['productividad_kg_ms_ha'], aliases=['Productividad:'], localize=True)
            ).add_to(m)
            folium.GeoJson(
                st.session_state.poligono_data.geometry.iloc[0],
                style_function=lambda x: {'fillColor': 'transparent', 'color': '#1d4ed8', 'weight': 3, 'dashArray': '5, 5'}
            ).add_to(m)
            colormap.add_to(m)
            folium_static(m, width=1000, height=600)
        except Exception as e:
            st.warning(f"No se pudo generar el mapa de sublotes: {str(e)}")

    # Calculadora interactiva
    with st.expander("📊 Calculadora de Equivalentes Vaca"):
        num_ev_input = st.number_input("Número de EV disponibles:", min_value=1.0, max_value=1000.0, value=50.0, step=1.0)
        dias_input = st.number_input("Días de permanencia deseada:", min_value=1, max_value=365, value=30, step=1)
        if st.button("Calcular días"):
            forrajero = forrajero_data['forrajero']
            dias_calc = forrajero.calcular_dias_permanencia(disp['forraje_aprovechable_kg_ms'], num_ev_input)
            st.success(f"**Resultado:** {num_ev_input:.0f} EV pueden pastar {dias_calc['dias_recomendados']} días")
            col1, col2, col3 = st.columns(3)
            with col1: st.metric("Días básicos", f"{dias_calc['dias_basico']:.1f}")
            with col2: st.metric("Días ajustados", f"{dias_calc['dias_ajustado']:.1f}")
            with col3: st.metric("Recomendados", dias_calc['dias_recomendados'])

def mostrar_comparacion():
    st.header("📈 Análisis Comparativo")
    if st.session_state.resultados is None:
        st.info("Ejecute el análisis primero.")
        return
    res = st.session_state.resultados
    if all(k in res for k in ['puntos_carbono', 'puntos_ndvi', 'puntos_ndwi', 'puntos_biodiversidad']):
        fig = Visualizaciones.crear_grafico_comparativo(
            res['puntos_carbono'], res['puntos_ndvi'], res['puntos_ndwi'], res['puntos_biodiversidad']
        )
        if fig:
            st.plotly_chart(fig, use_container_width=True)
    # Correlaciones
    st.subheader("🔗 Correlaciones")
    try:
        n = min(100, len(res['puntos_carbono']))
        carbono_vals = [p['carbono_ton_ha'] for p in res['puntos_carbono'][:n]]
        ndvi_vals = [p['ndvi'] for p in res['puntos_ndvi'][:n]]
        ndwi_vals = [p['ndwi'] for p in res['puntos_ndwi'][:n]]
        shannon_vals = [p['indice_shannon'] for p in res['puntos_biodiversidad'][:n]]
        corr1 = np.corrcoef(carbono_vals, ndvi_vals)[0,1] if len(carbono_vals)>1 else 0
        corr2 = np.corrcoef(carbono_vals, shannon_vals)[0,1] if len(carbono_vals)>1 else 0
        corr3 = np.corrcoef(ndvi_vals, shannon_vals)[0,1] if len(ndvi_vals)>1 else 0
        corr4 = np.corrcoef(ndwi_vals, shannon_vals)[0,1] if len(ndwi_vals)>1 else 0
        col1, col2, col3, col4 = st.columns(4)
        with col1: st.metric("C vs NDVI", f"{corr1:.3f}")
        with col2: st.metric("C vs Shannon", f"{corr2:.3f}")
        with col3: st.metric("NDVI vs Shannon", f"{corr3:.3f}")
        with col4: st.metric("NDWI vs Shannon", f"{corr4:.3f}")
    except Exception as e:
        st.warning(f"No se pudieron calcular correlaciones: {str(e)}")

def mostrar_prv():
    st.header("🐄 Pastoreo Racional Voisin (PRV)")
    if st.session_state.resultados is None or 'prv' not in st.session_state.resultados:
        st.info("Ejecute el análisis completo primero.")
        return

    res = st.session_state.resultados
    prv_data = res['prv']
    prv_modelo = prv_data['modelo']
    gdf_potreros = prv_data['gdf_potreros']
    plan_pastoreo = prv_data['plan_pastoreo']
    estados_prv = prv_data['estados_prv']
    resumen_prv = prv_data['resumen']

    if gdf_potreros is None or gdf_potreros.empty:
        st.warning("No se pudieron generar los potreros PRV. Pruebe con más puntos de muestreo o un área diferente.")
        return

    # === Configuración interactiva PRV ===
    with st.expander("⚙️ Ajustar parámetros PRV", expanded=False):
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            temperatura = st.number_input("Temperatura (°C)", min_value=5.0, max_value=40.0, value=float(prv_data.get('temperatura', 22)), step=1.0)
        with col2:
            precipitacion = st.number_input("Precipitación (mm/mes)", min_value=10.0, max_value=500.0, value=float(prv_data.get('precipitacion', 100)), step=10.0)
        with col3:
            num_potreros_ajuste = st.number_input("Número de potreros", min_value=4, max_value=80, value=int(prv_data.get('num_potreros', 20)), step=1)
        with col4:
            ocupacion_ajuste = st.number_input("Días de ocupación", min_value=1, max_value=5, value=int(prv_data.get('ocupacion_dias', 1)), step=1)

        if st.button("🔄 Recalcular PRV", use_container_width=True):
            descanso_nuevo = prv_modelo.calcular_periodo_descanso(temperatura, precipitacion)
            gdf_potreros_nuevo = prv_modelo.generar_potreros_desde_cuadricula(
                res.get('gdf_cuadricula'), num_potreros=num_potreros_ajuste
            )
            if gdf_potreros_nuevo is not None:
                plan_nuevo = prv_modelo.planificar_ciclo_prv(
                    num_potreros=len(gdf_potreros_nuevo),
                    descanso_dias=descanso_nuevo,
                    ocupacion_dias=ocupacion_ajuste,
                    num_ciclos=4
                )
                estados_nuevo = prv_modelo.calcular_estado_potreros(gdf_potreros_nuevo, plan_nuevo)
                resumen_nuevo = prv_modelo.resumen_ejecutivo(gdf_potreros_nuevo, estados_nuevo, plan_nuevo)
                recomendaciones_nuevo = prv_modelo.generar_recomendaciones_prv(
                    gdf_potreros_nuevo, estados_nuevo, temperatura, precipitacion
                )
                # Guardar en sesión
                st.session_state.resultados['prv']['gdf_potreros'] = gdf_potreros_nuevo
                st.session_state.resultados['prv']['plan_pastoreo'] = plan_nuevo
                st.session_state.resultados['prv']['estados_prv'] = estados_nuevo
                st.session_state.resultados['prv']['resumen'] = resumen_nuevo
                st.session_state.resultados['prv']['recomendaciones'] = recomendaciones_nuevo
                st.session_state.resultados['prv']['descanso_dias'] = descanso_nuevo
                st.session_state.resultados['prv']['ocupacion_dias'] = ocupacion_ajuste
                st.session_state.resultados['prv']['num_potreros'] = num_potreros_ajuste
                st.session_state.resultados['prv']['temperatura'] = temperatura
                st.session_state.resultados['prv']['precipitacion'] = precipitacion
                st.rerun()

    # === KPIs PRV ===
    if resumen_prv:
        st.subheader("📊 Resumen PRV")
        col1, col2, col3, col4, col5 = st.columns(5)
        with col1:
            st.metric("Potreros", resumen_prv.get('num_potreros', 0))
        with col2:
            st.metric("Área Total", f"{resumen_prv.get('area_total_ha', 0):.1f} ha")
        with col3:
            st.metric('Forraje Disp.', f"{resumen_prv.get('forraje_disponible_actual_kg_ms', 0)/1000:.1f} ton")
        with col4:
            st.metric("EV/30d", f"{resumen_prv.get('ev_soportables_30d', 0):.1f}")
        with col5:
            st.metric("Prod. media", f"{resumen_prv.get('productividad_prom_kg_ms_ha', 0):,.0f} kg/ha")

        col1, col2, col3, col4 = st.columns(4)
        with col1:
            desc = prv_data.get('descanso_dias', 0)
            st.metric("Descanso", f"{desc} días", delta="Ajustar cada 15d")
        with col2:
            oc = prv_data.get('ocupacion_dias', 1)
            st.metric("Ocupación", f"{oc} día(s)", delta="Ideal: 1 día")
        with col3:
            st.metric("En pastoreo", resumen_prv.get('potreros_en_pastoreo', 0))
        with col4:
            st.metric("Listos", resumen_prv.get('potreros_listos', 0), delta="Priorizar estos")

    # === Mapa de potreros PRV coloreado por productividad ===
    st.subheader("🗺️ Mapa de Potreros PRV")
    tab_map1, tab_map2 = st.tabs(["🎨 Por productividad", "🚦 Por estado"])
    sistema = SistemaMapas()

    with tab_map1:
        try:
            m = SistemaMapas.crear_mapa_con_base(gdf_potreros)

            min_prod = gdf_potreros['productividad_kg_ms_ha'].min()
            max_prod = gdf_potreros['productividad_kg_ms_ha'].max()

            colormap = LinearColormap(
                colors=['#8B0000', '#FF4500', '#FFD700', '#9ACD32', '#32CD32', '#006400'],
                vmin=min_prod, vmax=max_prod
            )
            colormap.caption = 'Productividad Forrajera (kg MS/ha) - PRV'

            # Destacar los potreros de mayor productividad (top 20%) con borde resaltado
            umbral_alto = gdf_potreros['productividad_kg_ms_ha'].quantile(0.8)

            for idx, row in gdf_potreros.iterrows():
                prod = row['productividad_kg_ms_ha']
                es_alta = prod >= umbral_alto
                folium.GeoJson(
                    row['geometry'],
                    style_function=lambda x, p=prod, alta=es_alta: {
                        'fillColor': colormap(p),
                        'color': '#FFD700' if alta else '#333',
                        'weight': 4 if alta else 1,
                        'fillOpacity': 0.75 if alta else 0.55,
                        'dashArray': None if alta else '3, 3',
                    },
                    tooltip=f"Potrero {row['potrero_id']}: {prod:,.0f} kg MS/ha{' ⭐ ALTA' if es_alta else ''}",
                    highlight_function=lambda x: {'weight': 5, 'color': '#FFD700', 'fillOpacity': 0.85}
                ).add_to(m)

            # Polígono del área total
            folium.GeoJson(
                st.session_state.poligono_data.geometry.iloc[0],
                style_function=lambda x: {'fillColor': 'transparent', 'color': '#1d4ed8', 'weight': 3, 'dashArray': '5, 5'}
            ).add_to(m)

            colormap.add_to(m)
            folium_static(m, width=1000, height=550)

            st.info("⭐ Los potreros con borde dorado son los de **mayor productividad** (top 20%). Priorícelos en la rotación.")

        except Exception as e:
            st.warning(f"No se pudo generar el mapa PRV: {str(e)}")

    with tab_map2:
        try:
            if not estados_prv.empty:
                estados_map = estados_prv.set_index('potrero_id')['estado'].to_dict()
                color_estado = {
                    'pastoreo': '#EF4444',
                    'descanso': '#3B82F6',
                    'listo': '#10B981',
                    'por_pastorear': '#F59E0B',
                    'sin_planificar': '#9CA3AF',
                }
                m2 = SistemaMapas.crear_mapa_con_base(gdf_potreros)

                for idx, row in gdf_potreros.iterrows():
                    estado = estados_map.get(row['potrero_id'], 'sin_planificar')
                    color = color_estado.get(estado, '#9CA3AF')
                    folium.GeoJson(
                        row['geometry'],
                        style_function=lambda x, c=color: {
                            'fillColor': c,
                            'color': '#333',
                            'weight': 1,
                            'fillOpacity': 0.65,
                        },
                        tooltip=f"Potrero {row['potrero_id']}: {estado.title()}",
                    ).add_to(m2)

                # Leyenda manual
                leyenda_html = """
                <div style="position:fixed; bottom:30px; left:30px; z-index:9999; background:white; padding:10px; border-radius:8px; box-shadow:0 0 10px rgba(0,0,0,0.3); font-size:12px;">
                    <b>Estado Potreros</b><br>
                    <span style='color:#EF4444;'>●</span> Pastoreo<br>
                    <span style='color:#3B82F6;'>●</span> Descanso<br>
                    <span style='color:#10B981;'>●</span> Listo<br>
                    <span style='color:#F59E0B;'>●</span> Por pastorear
                </div>
                """
                m2.get_root().html.add_child(folium.Element(leyenda_html))
                folium_static(m2, width=1000, height=550)

        except Exception as e:
            st.warning(f"No se pudo generar el mapa de estados: {str(e)}")

    # === Plan de pastoreo (calendario) ===
    if not plan_pastoreo.empty:
        st.subheader("📅 Calendario de Pastoreo PRV")
        df_plan = plan_pastoreo.copy()
        df_plan['fecha_entrada'] = df_plan['fecha_entrada'].dt.strftime('%d/%m/%Y')
        df_plan['fecha_salida'] = df_plan['fecha_salida'].dt.strftime('%d/%m/%Y')
        df_plan['fecha_proxima_entrada'] = df_plan['fecha_proxima_entrada'].dt.strftime('%d/%m/%Y')
        df_plan['potrero_id'] = df_plan['potrero_id'].astype(int)
        df_plan['ciclo'] = df_plan['ciclo'].astype(int)
        st.dataframe(df_plan, use_container_width=True, hide_index=True)

        # Gráfico de Gantt del plan de pastoreo
        fig_gantt = go.Figure()
        plan_plot = plan_pastoreo.copy()
        colores_potrero = [
            '#FF6B6B', '#4ECDC4', '#45B7D1', '#96CEB4', '#FFEAA7',
            '#DDA0DD', '#98D8C8', '#F7DC6F', '#BB8FCE', '#85C1E9',
            '#F0B27A', '#82E0AA', '#F1948A', '#85929E', '#73C6B6',
            '#E59866', '#AED6F1', '#D5F5E3', '#FADBD8', '#F9E79F',
        ]
        for potrero_id in sorted(plan_plot['potrero_id'].unique()):
            df_p = plan_plot[plan_plot['potrero_id'] == potrero_id]
            color = colores_potrero[(potrero_id - 1) % len(colores_potrero)]
            for _, row_p in df_p.iterrows():
                fig_gantt.add_trace(go.Scatter(
                    x=[row_p['fecha_entrada'], row_p['fecha_salida']],
                    y=[f"Potrero {potrero_id}", f"Potrero {potrero_id}"],
                    mode='lines+markers',
                    line=dict(color=color, width=12),
                    marker=dict(size=8, color=color),
                    name=f"Potrero {potrero_id}",
                    showlegend=(_ == 0),
                    hovertext=f"Ocupación: {row_p['dias_ocupacion']}d",
                ))

        fig_gantt.update_layout(
            title='Cronograma de Pastoreo - PRV',
            xaxis_title='Fecha',
            yaxis_title='Potrero',
            height=400 + len(plan_plot['potrero_id'].unique()) * 15,
            hovermode='closest',
            showlegend=False,
            xaxis=dict(
                rangeselector=dict(
                    buttons=list([
                        dict(count=1, label='1m', step='month', stepmode='backward'),
                        dict(count=3, label='3m', step='month', stepmode='backward'),
                        dict(step='all')
                    ])
                ),
                rangeslider=dict(visible=True),
                type='date',
            ),
        )
        st.plotly_chart(fig_gantt, use_container_width=True)

    # === Estado actual de potreros ===
    if not estados_prv.empty:
        st.subheader("🚦 Estado Actual de Potreros")
        df_estados = estados_prv.copy()
        df_estados['potrero_id'] = df_estados['potrero_id'].astype(int)

        # Color coding
        def color_estado_row(val):
            if val == 'pastoreo': return 'background-color: #FECACA'
            elif val == 'descanso': return 'background-color: #DBEAFE'
            elif val == 'listo': return 'background-color: #D1FAE5'
            elif val == 'por_pastorear': return 'background-color: #FEF3C7'
            return ''

        if hasattr(df_estados.style, 'map'):
            styled = df_estados.style.map(color_estado_row, subset=['estado'])
        else:
            styled = df_estados.style.applymap(color_estado_row, subset=['estado'])
        st.dataframe(
            styled,
            use_container_width=True,
            hide_index=True,
        )

    # === Curva de acumulación de forraje ===
    st.subheader("📈 Curva de Acumulación de Forraje")
    descanso_dias = prv_data.get('descanso_dias', 40)
    prod_prom_ha = resumen_prv.get('productividad_prom_kg_ms_ha', 3000)
    curva = prv_modelo.a_forraje_acumulado_curva(descanso_dias, prod_prom_ha)

    fig_curva = go.Figure()
    fig_curva.add_trace(go.Scatter(
        x=list(range(1, descanso_dias + 1)),
        y=curva,
        mode='lines',
        name='Acumulación',
        line=dict(color='#10B981', width=3),
        fill='tozeroy',
        fillcolor='rgba(16, 185, 129, 0.2)',
    ))
    # Marcar los días clave
    fig_curva.add_vline(x=descanso_dias, line_dash='dash', line_color='#EF4444',
                        annotation_text=f'Descanso: {descanso_dias}d',
                        annotation_position='top right')
    fig_curva.add_vline(x=max(1, descanso_dias // 2), line_dash='dot', line_color='#F59E0B',
                        annotation_text='Rebrote activo',
                        annotation_position='top left')
    fig_curva.update_layout(
        title=f'Acumulación Estimada de Forraje durante {descanso_dias} días de descanso',
        xaxis_title='Días de descanso',
        yaxis_title='Forraje acumulado (kg MS/ha)',
        height=400,
        hovermode='x unified',
    )
    st.plotly_chart(fig_curva, use_container_width=True)

    # === Recomendaciones PRV ===
    st.subheader("📋 Recomendaciones PRV")
    st.markdown(prv_data.get('recomendaciones', 'No hay recomendaciones disponibles.'))


def mostrar_asistente_prv():
    st.header("🤖 Asistente IA — Transición a PRV Regenerativo")
    st.markdown(
        '<div style="background:rgba(16,185,129,0.08);border:1px solid rgba(16,185,129,0.15);'
        'border-radius:12px;padding:1rem;margin-bottom:1.5rem;">'
        'Este asistente te ayuda a diseñar un plan personalizado de transición desde tu sistema actual'
        ' hacia el <strong>Pastoreo Racional Voisin (PRV)</strong> con enfoque de <strong>ganadería regenerativa</strong>.'
        ' Completá los datos de tu establecimiento y recibí un plan paso a paso.</div>',
        unsafe_allow_html=True
    )

    # Formulario
    with st.expander("📝 Datos del establecimiento", expanded=True):
        col1, col2 = st.columns(2)
        with col1:
            tipo_ganado = st.selectbox("Tipo de ganado", ["bovino", "ovino", "caprino", "bufalino", "mixto"], index=0)
            cabeza = st.number_input("Cantidad de animales", min_value=1, max_value=100000, value=100, step=10)
            area_ha = st.number_input("Superficie total (ha)", min_value=0.5, max_value=100000.0, value=100.0, step=10.0)
            sistema_actual = st.selectbox("Sistema actual de pastoreo", [
                "pastoreo continuo", "pastoreo rotativo simple (3-4 potreros)",
                "pastoreo rotativo intensivo (5-8 potreros)", "estabulado / confinamiento",
                "pastoreo en franjas", "sin pastoreo (solo heno/silo)"
            ], index=0)
        with col2:
            condicion_suelo = st.selectbox("Condición del suelo", ["degradado", "regular", "bueno", "excelente"], index=1)
            objetivo = st.selectbox("Objetivo principal", ["carne", "leche", "mixto (carne y leche)", "regeneración / carbono", "cria / recria"], index=0)
            tiene_agua = st.checkbox("Dispongo de agua en todos los potreros", value=True)
            tiene_divisiones = st.checkbox("Ya tengo divisiones / potreros", value=False)
            if tiene_divisiones:
                descanso_actual = st.slider("Días de descanso actuales", 0, 120, 30)
                ocupacion_actual = st.slider("Días de ocupación actuales", 1, 30, 7)
            else:
                descanso_actual = 0
                ocupacion_actual = 0

        prod_forraje_est = st.number_input(
            "Productividad forrajera estimada (kg MS/ha) — dejar en 0 si no se sabe",
            min_value=0, max_value=50000, value=0, step=500
        )
        ecosistema = st.text_input("Ecosistema / región", placeholder="ej. Pampa, Chaqueño, Espinal, etc.")

    if st.button("🤖 Generar plan de transición PRV", use_container_width=True, type="primary"):
        with st.spinner("El asistente IA está preparando tu plan personalizado..."):
            params = {
                "tipo_ganado": tipo_ganado,
                "cabeza": cabeza,
                "area_ha": area_ha,
                "sistema_actual": sistema_actual,
                "descanso_actual_dias": descanso_actual,
                "ocupacion_actual_dias": ocupacion_actual,
                "tiene_agua": tiene_agua,
                "tiene_divisiones": tiene_divisiones,
                "condicion_suelo": condicion_suelo,
                "objetivo_principal": objetivo,
                "ecosistema": ecosistema if ecosistema else "no especificado",
                "productividad_forraje": prod_forraje_est if prod_forraje_est > 0 else (st.session_state.resultados.get('analisis_forrajero', {}).get('productividad_prom_kg_ms_ha', 0) if st.session_state.resultados else 0),
            }
            plan = generar_plan_transicion_prv(params)

        st.divider()
        st.markdown("### 📋 Plan de Transición PRV")
        st.markdown(plan)

        st.divider()
        st.markdown(
            '<div style="background:rgba(59,130,246,0.06);border:1px solid rgba(59,130,246,0.1);'
            'border-radius:10px;padding:0.75rem 1rem;font-size:0.85rem;">'
            '💡 <strong>Sugerencia:</strong> Compartí este plan con tu asesor técnico y ajustalo'
            ' según las condiciones reales de tu campo. El PRV se construye observando y adaptando.</div>',
            unsafe_allow_html=True
        )

    # Si el análisis ya se ejecutó, mostrar info complementaria
    if st.session_state.resultados and 'prv' in st.session_state.resultados:
        prv_data = st.session_state.resultados['prv']
        with st.expander("📊 Datos PRV del análisis actual", expanded=False):
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("Potreros generados", prv_data.get('resumen', {}).get('num_potreros', 0))
            with col2:
                st.metric("Productividad media", f"{prv_data.get('resumen', {}).get('productividad_prom_kg_ms_ha', 0):,.0f} kg/ha")
            with col3:
                st.metric("EV/30d", prv_data.get('resumen', {}).get('ev_soportables_30d', 0))
            st.caption("Estos valores pueden usarse como referencia para el plan de transición.")


def mostrar_informe():
    st.header("📥 Informe Completo")
    if st.session_state.resultados is None or st.session_state.poligono_data is None:
        st.info("Ejecute el análisis primero.")
        return

    st.markdown("### Generar informe con todos los análisis")
    sistema = SistemaMapas()
    generador = GeneradorReportes(st.session_state.resultados, st.session_state.poligono_data, sistema)

    col1, col2, col3, col4 = st.columns(4)
    with col1:
        if REPORTPDF_AVAILABLE:
            if st.button("📄 Generar PDF", use_container_width=True):
                pdf = generador.generar_pdf()
                if pdf:
                    st.download_button("⬇️ Descargar PDF", pdf, f"informe_{datetime.now().strftime('%Y%m%d_%H%M')}.pdf", "application/pdf")
    with col2:
        if REPORTDOCX_AVAILABLE:
            if st.button("📘 Generar DOCX", use_container_width=True):
                docx = generador.generar_docx()
                if docx:
                    st.download_button("⬇️ Descargar DOCX", docx, f"informe_{datetime.now().strftime('%Y%m%d_%H%M')}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    with col3:
        # Verificar si la IA está disponible (cliente y API key)
        if GROQ_API_KEY is not None and groq_client is not None:
            if st.button("🤖 Generar Informe con IA (Groq)", use_container_width=True):
                with st.spinner("Generando informe con IA (Groq)..."):
                    reporte_ia = generar_reporte_ia(st.session_state.resultados, st.session_state.poligono_data, sistema)
                    if reporte_ia:
                        st.download_button("⬇️ Descargar Informe IA", reporte_ia, f"informe_IA_{datetime.now().strftime('%Y%m%d_%H%M')}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        else:
            st.info("🤖 IA no disponible (falta API key de Groq o cliente no configurado)")
    with col4:
        if st.button("🌍 Generar GeoJSON", use_container_width=True):
            geojson = generador.generar_geojson()
            if geojson:
                st.download_button("⬇️ Descargar GeoJSON", geojson, f"area_{datetime.now().strftime('%Y%m%d_%H%M')}.geojson", "application/geo+json")

# ===============================
# MAIN
# ===============================
def main():
    if 'gee_authenticated' not in st.session_state:
        st.session_state.gee_authenticated = False
        st.session_state.gee_project = ''
        if GEE_AVAILABLE:
            inicializar_gee()
    if 'poligono_data' not in st.session_state:
        st.session_state.poligono_data = None
    if 'resultados' not in st.session_state:
        st.session_state.resultados = None
    if 'mapa' not in st.session_state:
        st.session_state.mapa = None
    # Inicializar modelo seleccionado por defecto
    if 'selected_model' not in st.session_state:
        st.session_state.selected_model = available_models[0] if available_models else "llama3-70b-8192"

    st.markdown(
        '<div style="display:flex;align-items:center;gap:1rem;margin:0 0 0.25rem 0;">'
        '<div style="font-size:2.2rem;">🌎</div>'
        '<div><div style="font-size:1.6rem;font-weight:700;background:linear-gradient(135deg,#60a5fa,#a78bfa);'
        '-webkit-background-clip:text;-webkit-text-fill-color:transparent;">'
        'Sistema Satelital de Análisis Ambiental Integral</div>'
        '<div style="color:#64748b;font-size:0.85rem;font-weight:400;">'
        'Carbono · Biodiversidad · Análisis Forrajero · Pastoreo Racional Voisin</div></div></div>',
        unsafe_allow_html=True
    )

    with st.sidebar:
        st.markdown(
            '<div style="display:flex;align-items:center;gap:0.5rem;margin-bottom:1.5rem;">'
            '<span style="font-size:1.5rem;">🌿</span>'
            '<span style="font-size:1.1rem;font-weight:600;color:#f1f5f9;">Panel de Control</span></div>',
            unsafe_allow_html=True
        )

        st.markdown('<p style="color:#64748b;font-size:0.75rem;text-transform:uppercase;letter-spacing:0.05em;margin-bottom:0.25rem;">📁 Carga de Datos</p>', unsafe_allow_html=True)
        if GEE_AVAILABLE and st.session_state.gee_authenticated:
            st.markdown('<div style="background:rgba(16,185,129,0.1);border:1px solid rgba(16,185,129,0.2);border-radius:8px;padding:0.4rem 0.75rem;font-size:0.75rem;color:#6ee7b7;">✅ GEE Conectado</div>', unsafe_allow_html=True)
        uploaded_file = st.file_uploader("Cargar polígono (KML, GeoJSON, SHP, KMZ)", type=['kml', 'geojson', 'zip', 'kmz'],
                                         label_visibility="collapsed")
        if uploaded_file:
            with st.spinner("Procesando archivo..."):
                gdf = cargar_archivo_parcela(uploaded_file)
                if gdf is not None:
                    st.session_state.poligono_data = gdf
                    area_ha = calcular_superficie(gdf)
                    st.markdown(f'<div style="background:rgba(59,130,246,0.1);border:1px solid rgba(59,130,246,0.15);border-radius:8px;padding:0.4rem 0.75rem;font-size:0.8rem;color:#93c5fd;">📍 {area_ha:,.1f} ha</div>', unsafe_allow_html=True)
                    sistema = SistemaMapas()
                    st.session_state.mapa = sistema.crear_mapa_area(gdf)

        if st.session_state.poligono_data is not None:
            st.markdown('<hr style="border-color:rgba(255,255,255,0.05);margin:1rem 0;">', unsafe_allow_html=True)
            st.markdown('<p style="color:#64748b;font-size:0.75rem;text-transform:uppercase;letter-spacing:0.05em;margin-bottom:0.25rem;">⚙️ Configuración</p>', unsafe_allow_html=True)
            ecosistemas = [
                'amazonia', 'choco', 'andes', 'pampa', 'seco', 
                'cultivo', 'vid', 'agricola',
                'monte', 'espinal', 'yungas', 'chaqueño', 'patagonico', 'paranaense'
            ]
            tipo_ecosistema = st.selectbox("Ecosistema", ecosistemas, label_visibility="collapsed")
            num_puntos = st.slider("Puntos de muestreo", 10, 200, 50)
            usar_gee = False
            if GEE_AVAILABLE and st.session_state.gee_authenticated:
                usar_gee = st.checkbox("Usar datos reales de GEE")
            
            if available_models:
                st.session_state.selected_model = st.selectbox(
                    "Modelo IA",
                    available_models,
                    index=available_models.index(st.session_state.selected_model) if st.session_state.selected_model in available_models else 0,
                    label_visibility="collapsed",
                )
            
            st.markdown('<br>', unsafe_allow_html=True)
            if st.button("🚀 Ejecutar Análisis Completo", type="primary", use_container_width=True):
                with st.spinner("Analizando..."):
                    resultados = ejecutar_analisis_completo(st.session_state.poligono_data, tipo_ecosistema, num_puntos, usar_gee)
                    if resultados:
                        st.session_state.resultados = resultados
                        st.markdown('<div style="background:rgba(16,185,129,0.1);border:1px solid rgba(16,185,129,0.2);border-radius:8px;padding:0.5rem 0.75rem;font-size:0.85rem;color:#6ee7b7;text-align:center;">✅ Análisis completado</div>', unsafe_allow_html=True)

    if st.session_state.poligono_data is None:
        st.markdown(
            '<div style="display:flex;flex-direction:column;align-items:center;justify-content:center;'
            'min-height:50vh;text-align:center;gap:1rem;">'
            '<div style="font-size:4rem;opacity:0.6;">🌿</div>'
            '<div style="font-size:1.3rem;font-weight:500;color:#94a3b8;">Cargue un polígono en el panel lateral</div>'
            '<div style="color:#475569;font-size:0.85rem;max-width:380px;">'
            'Formatos soportados: KML, KMZ, GeoJSON o Shapefile (ZIP)</div></div>',
            unsafe_allow_html=True
        )
        with st.expander("📋 Información del Sistema"):
            st.markdown("""
            - 🌳 **Carbono** — Metodología Verra VCS
            - 🦋 **Biodiversidad** — Índice de Shannon
            - 📈 **Índices espectrales** — NDVI, NDWI, NDRE, MSAVI, EVI
            - 🐮 **Forrajero** — Productividad, EV, rotación
            - 🐄 **Pastoreo Racional Voisin** — Potreros, ciclos, calendario
            - 🗺️ **Mapas** — Interpolación KNN, calor continuo
            - 📊 **Dashboard** — KPIs, gráficos ejecutivos
            - 📄 **Informes** — PDF, DOCX, GeoJSON y con IA (Groq)
            - 🌍 **Ecosistemas argentinos** — monte, espinal, yungas, chaqueño, patagonico, paranaense
            """)
    else:
        tabs = st.tabs(["🗺️ Mapas", "📊 Dashboard", "🌳 Carbono", "🦋 Biodiversidad", "🐮 Forrajero", "🐄 PRV Voisin", "📈 Comparación", "📥 Informe", "🤖 Asistente PRV"])
        with tabs[0]: mostrar_mapas_calor()
        with tabs[1]: mostrar_dashboard()
        with tabs[2]: mostrar_carbono()
        with tabs[3]: mostrar_biodiversidad()
        with tabs[4]: mostrar_analisis_forrajero()
        with tabs[5]: mostrar_prv()
        with tabs[6]: mostrar_comparacion()
        with tabs[7]: mostrar_informe()
        with tabs[8]: mostrar_asistente_prv()

if __name__ == "__main__":
    main()
